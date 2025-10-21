# modules/ai/chat.py
# AI Chat Integration Helper Functions for Syntax Prime V2
# Provides integration functions for router.py - NO ENDPOINTS HERE
# Clean helper module with all integration logic for weather, prayer times, scraper, etc.
# Date: 9/26/25 - Converted to helper module only
# Date: 9/27/25 - Added Google Trends integration and Prayer Notifications
# Date: 9/28/25 - Added Voice Synthesis and Image Generation detection

#-- Section 1: Core Imports - 9/26/25
import os
import uuid
import json
import asyncio
from datetime import datetime
import pytz
from typing import Dict, List, Any, Optional
from pathlib import Path
import logging
import httpx
import os
database_url = os.getenv('DATABASE_URL')

# File processing imports
from PIL import Image, ImageOps
import pdfplumber
import pandas as pd
import magic
import cv2
import numpy as np
from io import BytesIO
import base64

# Word document processing
from docx import Document

# Excel file processing
import openpyxl
from openpyxl.utils import get_column_letter

# Python file processing
import ast
import tokenize
from io import StringIO

# Add to the existing imports section:
from .pattern_fatigue import get_pattern_fatigue_tracker, handle_duplicate_complaint, handle_time_joke_complaint

# Google Trends integration imports - 9/27/25
from ..integrations.google_trends.opportunity_detector import OpportunityDetector
from ..integrations.google_trends.opportunity_training import OpportunityTraining
from ..integrations.google_trends.integration_info import check_module_health

logger = logging.getLogger(__name__)

# File upload configuration
UPLOAD_DIR = Path("/home/app/uploads/chat_files")
MAX_FILE_SIZE = 25 * 1024 * 1024  # 10MB
ALLOWED_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.pdf', '.txt', '.md', '.csv', '.doc', '.docx', '.xls', '.xlsx', '.py'}

def ensure_upload_dir():
    """Ensure upload directory exists with proper error handling"""
    try:
        UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
        return True
    except PermissionError as e:
        logger.error(f"Cannot create upload directory: {e}")
        return False
    except Exception as e:
        logger.error(f"Unexpected error creating upload directory: {e}")
        return False

_upload_dir_created = False

def get_upload_dir():
    """Get upload directory, creating it if necessary"""
    global _upload_dir_created
    if not _upload_dir_created:
        _upload_dir_created = ensure_upload_dir()
    return UPLOAD_DIR if _upload_dir_created else None

#-- Section 2: DateTime Helper Functions - 9/26/25
def get_current_datetime_context() -> dict:
    """
    Get comprehensive current date/time context for AI personalities
    Returns properly formatted date/time info in multiple formats
    """
    now = datetime.now()
    
    # Get user's timezone (defaulting to EST if not available)
    try:
        user_timezone = pytz.timezone('America/New_York')  # Default to Eastern
        now_user_tz = now.astimezone(user_timezone)
    except:
        now_user_tz = now
    
    return {
        "current_date": now_user_tz.strftime("%Y-%m-%d"),  # 2025-09-26
        "current_time_24h": now_user_tz.strftime("%H:%M"),  # 15:19
        "current_time_12h": now_user_tz.strftime("%I:%M %p"),  # 3:19 PM
        "day_of_week": now_user_tz.strftime("%A"),  # Friday
        "month_name": now_user_tz.strftime("%B"),  # September
        "full_datetime": now_user_tz.strftime("%A, %B %d, %Y at %H:%M"),  # Friday, September 26, 2025 at 15:19
        "timezone": str(now_user_tz.tzinfo),
        "iso_timestamp": now_user_tz.isoformat(),
        "unix_timestamp": int(now_user_tz.timestamp())
    }

#-- Section 3: Weather Integration Functions - 9/26/25
def detect_weather_request(message: str) -> tuple[bool, str]:
    """Detect weather-related requests
    
    Returns:
        tuple: (is_weather_request, request_type)
        request_type: 'current' or 'forecast'
    
    """
    weather_keywords = [
        "weather", "temperature", "forecast", "rain", "snow", "sunny",
        "cloudy", "storm", "wind", "humidity", "barometric pressure",
        "headache weather", "pressure change", "uv index"
    ]
    message_lower = message.lower()
    
    # Check if it's a weather request
    is_weather = any(keyword in message_lower for keyword in weather_keywords)
    
    if not is_weather:
        return (False, None)
    
    # Determine if it's a forecast request
    forecast_keywords = ["forecast", "tomorrow", "next", "upcoming", "week", "days"]
    is_forecast = any(keyword in message_lower for keyword in forecast_keywords)
    
    request_type = 'forecast' if is_forecast else 'current'
    
    return (True, request_type)

async def get_weather_forecast_for_user(user_id: str, location: str = None) -> Dict:
    """Get weather forecast for the user"""
    try:
        from modules.integrations.weather.tomorrow_client import TomorrowClient
        
        client = TomorrowClient()
        # Use IP-based location detection (reuse prayer times infrastructure)
        if not location:
            try:
                from modules.integrations.prayer_times.location_detector import detect_user_location
                location_data = await detect_user_location()
                location = f"{location_data['latitude']},{location_data['longitude']}"
                logger.info(f"🌍 Weather using IP location: {location_data['city']}, {location_data['region']}")
            except Exception as e:
                logger.warning(f"IP location failed, using fallback: {e}")
                location = "38.8606,-77.2287"  # Fallback to Fairfax, VA
        
        # Get forecast data
        forecast_data = await client.get_weather_forecast(location, days=5)
        
        if not forecast_data:
            return {
                "success": False,
                "error": "No forecast data available"
            }
        
        # Format forecast data
        forecast_days = []
        
        # Weather code mappings
        WEATHER_CODES = {
            1000: "Clear, Sunny", 1100: "Mostly Clear", 1101: "Partly Cloudy",
            1102: "Mostly Cloudy", 1001: "Cloudy", 2000: "Fog", 4000: "Drizzle",
            4001: "Rain", 4200: "Light Rain", 4201: "Heavy Rain", 5000: "Snow",
            5100: "Light Snow", 5101: "Heavy Snow", 8000: "Thunderstorm"
        }
        
        for day in forecast_data:
            values = day.get('values', {})
            timestamp = day.get('time', '')
            
            # Parse date
            from datetime import datetime
            date_obj = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
            day_name = date_obj.strftime('%A, %B %d')
            
            # Convert temperatures from Celsius to Fahrenheit
            temp_min_c = values.get('temperatureMin', 0)
            temp_max_c = values.get('temperatureMax', 0)
            temp_min_f = temp_min_c * 9/5 + 32
            temp_max_f = temp_max_c * 9/5 + 32
            
            # Get weather description
            weather_code = values.get('weatherCodeMax', 1000)
            condition = WEATHER_CODES.get(weather_code, "Unknown")
            
            forecast_days.append({
                'day_name': day_name,
                'temp_high': f"{temp_max_f:.1f}",
                'temp_low': f"{temp_min_f:.1f}",
                'condition': condition,
                'precipitation_probability': values.get('precipitationProbabilityAvg', 0),
                'humidity': values.get('humidityAvg', 0),
                'wind_speed': values.get('windSpeedAvg', 0),
                'uv_index': values.get('uvIndexMax', 0)
            })
        
        return {
            "success": True,
            "data": {
                "location": location,
                "forecast_days": forecast_days
            }
        }
        
    except Exception as e:
        logger.error(f"Forecast error: {e}")
        return {
            "success": False,
            "error": str(e)
        }

async def get_weather_for_user(user_id: str, location: str = None) -> Dict:
    """Get current weather data for the user"""
    try:
        from modules.integrations.weather.tomorrow_client import TomorrowClient
        
        client = TomorrowClient()
        # Use IP-based location detection (reuse prayer times infrastructure)
        if not location:
            try:
                from modules.integrations.prayer_times.location_detector import detect_user_location
                location_data = await detect_user_location()
                location = f"{location_data['latitude']},{location_data['longitude']}"
                logger.info(f"🌍 Weather using IP location: {location_data['city']}, {location_data['region']}")
            except Exception as e:
                logger.warning(f"IP location failed, using fallback: {e}")
                location = "38.8606,-77.2287"  # Fallback to Fairfax, VA
        
        # Get current weather data
        weather_data = await client.get_current_weather(location)
        
        if not weather_data:
            return {
                "success": False,
                "error": "No weather data available"
            }
        
        # Convert temperature from Celsius to Fahrenheit
        temp_c = weather_data.get('temperature', 0)
        temp_f = temp_c * 9/5 + 32
        
        # Format the response
        return {
            "success": True,
            "data": {
                "location": location,
                "temperature_f": f"{temp_f:.1f}",
                "humidity": weather_data.get('humidity', 0),
                "wind_speed": weather_data.get('windSpeed', 0),
                "pressure": weather_data.get('pressureSurfaceLevel', 0),
                "uv_index": weather_data.get('uvIndex', 0),
                "visibility": weather_data.get('visibility', 0)
            }
        }
        
    except Exception as e:
        logger.error(f"Current weather error: {e}")
        return {
            "success": False,
            "error": str(e)
        }

#-- Section 4: Bluesky Integration Functions - 9/26/25
#-- Section 4: Bluesky Integration Functions updated for posting - 9/30/25
import re
from typing import Tuple, Optional

def detect_bluesky_post_command(message: str) -> Tuple[bool, str, Optional[str], Optional[str]]:
    """
    Detect V1-style Bluesky posting commands
    
    Returns: (is_post_command, command_type, account, text)
    command_type: 'direct', 'write', 'improve', 'smart'
    """
    message_lower = message.lower().strip()
    
    # Pattern 1: bluesky post [account] "text"
    direct_pattern = r'bluesky post (\w+)\s+["\'](.+?)["\']'
    match = re.search(direct_pattern, message, re.IGNORECASE)
    if match:
        account = match.group(1)
        text = match.group(2)
        return True, 'direct', account, text
    
    # Pattern 2: write bluesky post about [topic] for [account]
    write_pattern = r'write bluesky post about (.+?) for (\w+)'
    match = re.search(write_pattern, message, re.IGNORECASE)
    if match:
        topic = match.group(1).strip()
        account = match.group(2).strip()
        return True, 'write', account, topic
    
    # Pattern 3: improve bluesky post: [text]
    improve_pattern = r'improve bluesky post:\s*(.+)'
    match = re.search(improve_pattern, message, re.IGNORECASE | re.DOTALL)
    if match:
        text = match.group(1).strip()
        return True, 'improve', None, text
    
    # Pattern 4: bluesky post smart "text"
    smart_pattern = r'bluesky post smart\s+["\'](.+?)["\']'
    match = re.search(smart_pattern, message, re.IGNORECASE)
    if match:
        text = match.group(1)
        return True, 'smart', None, text
    
    return False, '', None, None

def detect_bluesky_command(message: str) -> bool:
    """Detect ALL Bluesky commands including V1-style posting"""
    
    # Check for V1-style posting commands first
    is_post, _, _, _ = detect_bluesky_post_command(message)
    if is_post:
        return True
    
    # Original detection for management commands
    bluesky_keywords = [
        "bluesky scan", "bluesky opportunities", "bluesky accounts",
        "bluesky health", "bluesky status", "bluesky", "social media opportunities"
    ]
    message_lower = message.lower()
    return any(keyword in message_lower for keyword in bluesky_keywords)

async def process_bluesky_post_command(command_type: str, account: Optional[str],
                                      text: str, user_id: str) -> str:
    """Process V1-style Bluesky posting commands"""
    try:
        from ..integrations.bluesky.multi_account_client import get_bluesky_multi_client
        
        multi_client = get_bluesky_multi_client()
        
        # Account name mapping
        account_map = {
            'syntaxprime': 'personal',
            'personal': 'personal',
            'roseangel': 'rose_angel',
            'rose': 'rose_angel',
            'rose_angel': 'rose_angel',
            'bingetv': 'binge_tv',
            'tv': 'binge_tv',
            'binge_tv': 'binge_tv',
            'mealsfeelz': 'meals_feelz',
            'meals': 'meals_feelz',
            'meals_feelz': 'meals_feelz',
            'carl': 'damn_it_carl',
            'damnitcarl': 'damn_it_carl',
            'damn_it_carl': 'damn_it_carl'
        }
        
        # Command 1: Direct Post
        if command_type == 'direct':
            account_id = account_map.get(account.lower())
            
            if not account_id:
                return f"❌ Unknown account: {account}\n\nAvailable: syntaxprime, roseangel, bingetv, mealsfeelz, damnitcarl"
            
            if len(text) > 300:
                return f"❌ Post too long! ({len(text)}/300 characters)\n\nPlease shorten your post."
            
            result = await multi_client.create_post(account_id, text)
            
            if result['success']:
                account_info = multi_client.get_account_info(account_id)
                return f"""✅ Posted to {account_info['handle']}!

📝 "{text}"

Character count: {len(text)}/300"""
            else:
                return f"❌ Failed to post: {result.get('error', 'Unknown error')}"
        
        # Command 2: AI Write Post
        elif command_type == 'write':
            account_id = account_map.get(account.lower())
            
            if not account_id:
                return f"❌ Unknown account: {account}\n\nAvailable: syntaxprime, roseangel, bingetv, mealsfeelz, damnitcarl"
            
            account_info = multi_client.get_account_info(account_id)
            personality = account_info['personality']
            
            ai_post = await _generate_ai_post(text, personality, account_id)
            
            return f"""🤖 AI-Generated Post for {account_info['handle']}

📝 **Draft:**
"{ai_post}"

Character count: {len(ai_post)}/300

To post this, say:
`bluesky post {account} "{ai_post}"`

Or edit it first!"""
        
        # Command 3: Improve Post
        elif command_type == 'improve':
            improved_text = await _improve_post_text(text)
            
            return f"""✨ Improved Version:

"{improved_text}"

Character count: {len(improved_text)}/300

Original:
"{text}"

To post this, choose an account:
`bluesky post [account] "{improved_text}"`"""
        
        # Command 4: Smart Post
        elif command_type == 'smart':
            best_account = await _pick_best_account(text, multi_client)
            
            if len(text) > 300:
                return f"❌ Post too long! ({len(text)}/300 characters)\n\nPlease shorten your post."
            
            result = await multi_client.create_post(best_account, text)
            
            if result['success']:
                account_info = multi_client.get_account_info(best_account)
                return f"""🎯 Smart Post - Selected {account_info['handle']}

📝 "{text}"

Why this account: {_get_account_reasoning(best_account, text)}

Character count: {len(text)}/300"""
            else:
                return f"❌ Failed to post: {result.get('error', 'Unknown error')}"
        
    except Exception as e:
        return f"❌ Error: {str(e)}"

async def _generate_ai_post(topic: str, personality: str, account_id: str) -> str:
    """Generate AI post based on topic and personality"""
    
    if 'coding' in topic.lower() or 'web' in topic.lower() or 'python' in topic.lower():
        posts = {
            'syntaxprime': f"Hot take: {topic} is actually way simpler than everyone makes it sound. Here's why... 🧵",
            'professional': f"Key insights on {topic} for modern development practices. Worth considering for your next project.",
            'compassionate': f"Learning about {topic}? Remember: everyone starts somewhere. You've got this! 💚"
        }
    else:
        posts = {
            'syntaxprime': f"Real talk about {topic} - it's not what you think. (Spoiler: it's better) ✨",
            'professional': f"Important considerations regarding {topic}. Essential reading for professionals in the field.",
            'compassionate': f"Thinking about {topic} today. It's important to approach this with empathy and understanding. 🌟"
        }
    
    return posts.get(personality, posts['syntaxprime'])

async def _improve_post_text(original_text: str) -> str:
    """Improve post text while maintaining meaning"""
    
    improved = original_text.strip()
    
    # Add emoji if missing and appropriate
    if not any(char in improved for char in '😀😁😂🤣😃😄😅😆😉😊😋😎✨🚀💚🌟'):
        if 'great' in improved.lower() or 'awesome' in improved.lower():
            improved += " ✨"
        elif 'think' in improved.lower() or 'consider' in improved.lower():
            improved += " 🤔"
    
    # Ensure proper capitalization
    if improved and improved[0].islower():
        improved = improved[0].upper() + improved[1:]
    
    # Trim if too long
    if len(improved) > 280:
        improved = improved[:277] + "..."
    
    return improved

async def _pick_best_account(text: str, multi_client) -> str:
    """Analyze text content and pick best account"""
    
    text_lower = text.lower()
    
    if any(word in text_lower for word in ['code', 'coding', 'python', 'javascript', 'web', 'dev', 'tech']):
        return 'personal'
    
    elif any(word in text_lower for word in ['business', 'nonprofit', 'consulting', 'strategy', 'professional']):
        return 'rose_angel'
    
    elif any(word in text_lower for word in ['tv', 'show', 'movie', 'watch', 'streaming', 'netflix', 'series']):
        return 'binge_tv'
    
    elif any(word in text_lower for word in ['food', 'meal', 'recipe', 'cooking', 'hunger', 'charity', 'giving']):
        return 'meals_feelz'
    
    elif any(word in text_lower for word in ['creative', 'burnout', 'therapy', 'art', 'design']):
        return 'damn_it_carl'
    
    return 'personal'

def _get_account_reasoning(account_id: str, text: str) -> str:
    """Explain why this account was chosen"""
    
    reasons = {
        'personal': "Content matches your tech/coding expertise",
        'rose_angel': "Professional consulting/business content",
        'binge_tv': "Entertainment and streaming content",
        'meals_feelz': "Food, charity, or community focus",
        'damn_it_carl': "Creative and personal expression"
    }
    
    return reasons.get(account_id, "Best match for your content")

async def process_bluesky_command(message: str, user_id: str) -> str:
    """Process Bluesky-related commands"""
    try:
        # Check for V1-style posting commands FIRST
        is_post, cmd_type, account, text = detect_bluesky_post_command(message)
        if is_post:
            return await process_bluesky_post_command(cmd_type, account, text, user_id)
        
        # Original management commands
        from ..integrations.bluesky.multi_account_client import get_bluesky_multi_client
        from ..integrations.bluesky.engagement_analyzer import get_engagement_analyzer
        from ..integrations.bluesky.approval_system import get_approval_system
        from ..integrations.bluesky.notification_manager import get_notification_manager
        
        message_lower = message.lower()
        
        if 'scan' in message_lower:
            await trigger_background_scan(user_id)
            return """🔵 **Bluesky Account Scan Initiated**

✅ Scanning all configured accounts for engagement opportunities...
🔍 Analyzing posts from the last 24 hours
🤖 AI-powered engagement suggestions incoming
📝 Draft posts will be generated for approval

Results will be available in a few moments. Use `bluesky opportunities` to view suggestions."""
        
        elif 'opportunities' in message_lower or 'suggestion' in message_lower:
            approval_system = get_approval_system()
            pending_items = await approval_system.get_pending_approvals(limit=5)
            
            if not pending_items:
                return """🔵 **No Current Opportunities**

No engagement opportunities found at this time.
• Use `bluesky scan` to search for new opportunities
• Check back in a few hours for automatic updates"""
            
            response_parts = ["📊 **Current Bluesky Engagement Opportunities**\n"]
            
            for i, item in enumerate(pending_items, 1):
                account_info = item.get('account_info', {})
                opportunity = item.get('opportunity_analysis', {})
                
                response_parts.append(f"""**{i}. {account_info.get('username', 'Unknown Account')}**
📊 Engagement Score: {opportunity.get('engagement_potential', 0):.0%}
💡 Why: {opportunity.get('opportunity_reason', 'High engagement potential')}
⏰ Post Time: {opportunity.get('post_time', 'Recently')}

""")
            
            response_parts.append("Use the Bluesky dashboard to review and approve these opportunities.")
            return "\n".join(response_parts)
        
        elif 'accounts' in message_lower or 'status' in message_lower:
            multi_client = get_bluesky_multi_client()
            account_statuses = multi_client.get_all_accounts_status()
            configured_count = len([s for s in account_statuses.values() if s.get('configured', False)])
            
            response_parts = ["**Bluesky Account Status**\n"]
            
            for account_id, status in account_statuses.items():
                emoji = "✅" if status.get('authenticated', False) else "❌"
                username = status.get('username', f'Account {account_id}')
                last_scan = status.get('last_scan', 'Never')
                
                response_parts.append(f"{emoji} **{username}**")
                response_parts.append(f"   Last scan: {last_scan}")
                response_parts.append("")
            
            response_parts.append(f"**Summary:** {configured_count}/5 accounts configured and ready")
            return "\n".join(response_parts)
        
        else:
            account_statuses = get_bluesky_multi_client().get_all_account_status()
            return f"""🔵 **Bluesky Social Media Intelligence**

📱 **Configured Accounts:** {len(account_statuses)}/5
🤖 **Features:** Keyword intelligence, engagement suggestions, approval workflow
⏰ **Auto-Scan:** Every 3.5 hours across all accounts

**Available Commands:**
• `bluesky scan` - Force scan all accounts  
• `bluesky opportunities` - View engagement suggestions
• `bluesky accounts` - Check account status
• `bluesky post [account] "text"` - Post directly
• `write bluesky post about [topic] for [account]` - AI generates post
• `improve bluesky post: [text]` - AI improves your draft
• `bluesky post smart "text"` - AI picks best account

Ready to find your next great engagement opportunity?"""
    
    except Exception as e:
        logger.error(f"Bluesky command processing failed: {e}")
        return f"❌ **Bluesky Command Error:** {str(e)}\n\nTry `bluesky health` to check system status."

async def trigger_background_scan(user_id: str):
    """Trigger a background scan of all Bluesky accounts"""
    try:
        from ..integrations.bluesky.multi_account_client import get_bluesky_multi_client
        from ..integrations.bluesky.engagement_analyzer import get_engagement_analyzer
        from ..integrations.bluesky.approval_system import get_approval_system
        
        multi_client = get_bluesky_multi_client()
        engagement_analyzer = get_engagement_analyzer()
        approval_system = get_approval_system()
        
        await multi_client.authenticate_all_accounts()
        timelines = await multi_client.get_all_timelines()
        
        all_opportunities = []
        for account_id, timeline in timelines.items():
            account_config = multi_client.get_account_info(account_id)
            
            for post in timeline[:20]:
                analysis = await engagement_analyzer.analyze_post_for_account(
                    post, account_id, account_config
                )
                
                if analysis and analysis.get('engagement_potential', 0) >= 0.3:
                    draft_result = await approval_system.generate_draft_post(analysis, "reply")
                    
                    if draft_result['success']:
                        await approval_system.create_approval_item(
                            analysis, draft_result, analysis.get('priority_level', 'medium')
                        )
                        all_opportunities.append(analysis)
        
        logger.info(f"Background scan complete: {len(all_opportunities)} opportunities found")
        
    except Exception as e:
        logger.error(f"Background scan failed: {e}")

#-- Section 5: RSS Learning Functions - 9/26/25
def detect_rss_command(message: str) -> bool:
    """Detect if user is asking for RSS/marketing insights"""
    rss_keywords = [
        "marketing trends", "content ideas", "writing inspiration", "blog ideas",
        "social media trends", "seo trends", "marketing insights", "campaign ideas",
        "content strategy", "latest marketing", "marketing news",
        "industry trends", "rss insights", "marketing research", "content research", "writing help"
    ]
    message_lower = message.lower()
    return any(keyword in message_lower for keyword in rss_keywords)

def detect_writing_assistance_request(message: str) -> tuple[bool, str]:
    """Detect writing assistance requests and determine content type"""
    message_lower = message.lower()
    
    if any(term in message_lower for term in ['email', 'newsletter', 'email campaign', 'email marketing']):
        return True, 'email'
    elif any(term in message_lower for term in ['blog post', 'blog', 'article', 'write blog']):
        return True, 'blog'
    elif any(term in message_lower for term in ['social media', 'social post', 'tweet', 'linkedin post', 'instagram']):
        return True, 'social'
    elif any(term in message_lower for term in ['help me write', 'writing help', 'content ideas', 'what should i write']):
        return True, 'blog'
    
    return False, ''

async def get_rss_marketing_context(message: str, content_type: str = None) -> str:
    """Get marketing context from RSS learning system for AI writing assistance"""
    try:
        from ..integrations.rss_learning.marketing_insights import MarketingInsightsExtractor
        
        insights_extractor = MarketingInsightsExtractor()
        
        is_writing_request, detected_type = detect_writing_assistance_request(message)
        final_content_type = content_type or detected_type or 'blog'
        
        if is_writing_request or detect_rss_command(message):
            inspiration = await insights_extractor.get_writing_inspiration(
                content_type=final_content_type,
                topic=None,
                target_audience="digital marketers"
            )
            
            trends = await insights_extractor.get_latest_trends(limit=3)
            
            context_parts = [
                "CURRENT MARKETING INSIGHTS FROM RSS LEARNING SYSTEM:",
                ""
            ]
            
            if trends.get('trends_summary'):
                context_parts.extend([
                    "CURRENT TRENDS:",
                    trends['trends_summary'],
                    ""
                ])
            
            if inspiration.get('content_ideas'):
                context_parts.extend([
                    f"CONTENT IDEAS FOR {final_content_type.upper()}:",
                    *[f"• {idea}" for idea in inspiration['content_ideas'][:3]],
                    ""
                ])
            
            if inspiration.get('key_messages'):
                context_parts.extend([
                    "KEY MESSAGES TO CONSIDER:",
                    *[f"• {msg}" for msg in inspiration['key_messages'][:3]],
                    ""
                ])
            
            if trends.get('actionable_insights'):
                context_parts.extend([
                    "ACTIONABLE MARKETING INSIGHTS:",
                    *[f"• {insight}" for insight in trends['actionable_insights'][:3]],
                    ""
                ])
            
            if trends.get('trending_keywords'):
                context_parts.extend([
                    "TRENDING KEYWORDS:",
                    f"Consider incorporating: {', '.join(trends['trending_keywords'][:5])}",
                    ""
                ])
            
            if inspiration.get('call_to_action_ideas'):
                context_parts.extend([
                    "CALL-TO-ACTION IDEAS:",
                    *[f"• {cta}" for cta in inspiration['call_to_action_ideas'][:2]],
                    ""
                ])
            
            context_parts.extend([
                "Use these insights naturally in your response. Don't simply list them - integrate them into helpful, actionable advice for the user's specific request.",
                "Focus on current marketing best practices and trending approaches."
            ])
            
            return "\n".join(context_parts)
        
        elif detect_rss_command(message):
            trends = await insights_extractor.get_latest_trends(limit=5)
            
            context_parts = [
                "CURRENT MARKETING INSIGHTS:",
                "",
                "TRENDS SUMMARY:",
                trends.get('trends_summary', 'No current trends available'),
                ""
            ]
            
            if trends.get('trending_keywords'):
                context_parts.extend([
                    "TRENDING TOPICS:",
                    f"{', '.join(trends['trending_keywords'][:8])}",
                    ""
                ])
            
            context_parts.append("Provide insights based on these current marketing trends and data.")
            return "\n".join(context_parts)
        
        return ""
        
    except Exception as e:
        logger.error(f"RSS marketing context generation failed: {e}")
        return "RSS_CONTEXT_ERROR: Unable to retrieve current marketing insights. Proceed with general knowledge."

#-- Section 6: Marketing Scraper Functions - 9/26/25
def detect_scraper_command(message: str) -> bool:
    """Detect marketing scraper commands"""
    scraper_keywords = [
        "scrape", "scraper", "analyze website", "competitor analysis", "scrape url",
        "scrape site", "website analysis", "marketing analysis", "content analysis",
        "scrape history", "scrape insights", "scrape data"
    ]
    message_lower = message.lower()
    return any(keyword in message_lower for keyword in scraper_keywords)

def extract_url_from_message(message: str) -> str:
    """Extract URL from scrape command message"""
    import re
    
    # Look for URLs in the message
    url_pattern = r'https?://[^\s<>"{\}|\\^`\[\]]+'
    urls = re.findall(url_pattern, message)
    
    if urls:
        return urls[0]  # Return first URL found
    
    # If no full URL, look for domain patterns
    domain_pattern = r'(?:^|\s)([a-zA-Z0-9-]+\.(?:com|org|net|edu|gov|io|co\.uk))'
    domains = re.findall(domain_pattern, message)
    
    if domains:
        return f"https://{domains[0]}"
    
    return None

async def process_scraper_command(message: str, user_id: str) -> str:
    """Process marketing scraper commands"""
    try:
        from ..integrations.marketing_scraper.scraper_client import MarketingScraperClient
        from ..integrations.marketing_scraper.content_analyzer import ContentAnalyzer
        from ..integrations.marketing_scraper.database_manager import ScrapedContentDatabase
        
        message_lower = message.lower()
        
        if 'scrape history' in message_lower:
            # Get scrape history
            db = ScrapedContentDatabase()
            history = await db.get_user_scrape_history(user_id=user_id, limit=10)
            
            if not history:
                return """🔍 **Marketing Scraper History**

No scraping history found. Start analyzing competitors with:
• `scrape https://example.com` - Analyze any website
• `scrape insights` - Get analysis from previous scrapes"""
            
            response_parts = ["**Recent Scraping History**\n"]
            
            for i, item in enumerate(history, 1):
                domain = item.get('domain', 'Unknown')
                scraped_at = item.get('created_at', 'Unknown time')
                word_count = item.get('word_count', 0)
                
                response_parts.append(f"**{i}. {domain}**")
                response_parts.append(f"   📅 Scraped: {scraped_at}")
                response_parts.append(f"   📄 Words: {word_count}")
                response_parts.append("")
            
            response_parts.append("💡 Use `scrape insights` to get AI analysis of all scraped content.")
            return "\n".join(response_parts)
        
        elif 'scrape insights' in message_lower:
            # Get competitive insights from all scraped content
            db = ScrapedContentDatabase()
            
            # Search for recent content using empty topic to get all
            recent_content = await db.search_scraped_insights(user_id=user_id, topic="", limit=20)
            
            if not recent_content:
                return """🔍 **Marketing Scraper Insights**

No scraped content available for analysis. 

Start building your competitive intelligence with:
• `scrape https://competitor.com` - Analyze competitor sites
• `scrape https://industry-blog.com` - Analyze industry content"""
            
            # Generate competitive insights summary
            response_parts = [
                "🧠 **Competitive Intelligence Report**",
                f"📊 Based on {len(recent_content)} recently analyzed websites",
                ""
            ]
            
            # Show key insights from stored content
            for i, content in enumerate(recent_content[:5], 1):
                insights = content.get('key_insights', {})
                response_parts.append(f"**{i}. {content['domain']}**")
                if insights.get('value_proposition'):
                    response_parts.append(f"   💎 Value Prop: {insights['value_proposition'][:100]}...")
                if insights.get('content_strategy'):
                    response_parts.append(f"   📄 Strategy: {insights['content_strategy'][:100]}...")
                response_parts.append("")
            
            response_parts.append("📈 Use `scrape https://newsite.com` to add more competitive intelligence!")
            return "\n".join(response_parts)
        
        else:
            # Extract URL and scrape content
            url = extract_url_from_message(message)
            
            if not url:
                return """🔍 **Marketing Scraper Commands**

**Usage:**
• `scrape https://example.com` - Analyze any website for marketing insights
• `scrape history` - View your scraping history  
• `scrape insights` - Get competitive intelligence report

**Examples:**
• `scrape https://hubspot.com/blog` - Analyze HubSpot's content strategy
• `scrape https://competitor.com` - Competitive analysis
• `scrape https://industry-news.com` - Industry trend analysis

Ready to analyze your competition? 🕵️"""
            
            # Perform the scrape
            scraper = MarketingScraperClient()
            analyzer = ContentAnalyzer()
            db = ScrapedContentDatabase()
            
            try:
                # Scrape the website
                scraped_data = await scraper.scrape_website(url)
                
                if not scraped_data.get('scrape_status') == 'completed':
                    return f"""❌ **Scraping Failed**
                    
Unable to analyze {url}
Error: {scraped_data.get('error_message', 'Unknown error')}

Please verify the URL is accessible and try again."""
                
                # Analyze the content
                analysis = await analyzer.analyze_scraped_content(scraped_data)
                
                # Store in database
                await db.store_scraped_content(
                    user_id=user_id,
                    scraped_data=scraped_data,
                    analysis_results=analysis
                )
                
                # Generate response
                domain = scraped_data.get('domain', url)
                word_count = scraped_data.get('word_count', 0)
                
                response_parts = [
                    f"✅ **Successfully Analyzed: {domain}**",
                    f"📄 Content extracted: {word_count:,} words",
                    ""
                ]
                
                if analysis.get('competitive_insights'):
                    insights = analysis.get('competitive_insights', {})
                    if insights.get('value_proposition'):
                        response_parts.extend([
                            "**💎 Value Proposition:**",
                            f"• {insights['value_proposition'][:200]}...",
                            ""
                        ])
                
                if analysis.get('marketing_angles'):
                    marketing = analysis.get('marketing_angles', {})
                    if marketing.get('content_strategy'):
                        response_parts.extend([
                            "**📄 Content Strategy:**",
                            f"• {marketing['content_strategy'][:200]}...",
                            ""
                        ])
                
                if analysis.get('cta_analysis'):
                    cta = analysis.get('cta_analysis', {})
                    if cta.get('cta_placement_strategy'):
                        response_parts.extend([
                            "**🔥 CTA Strategy:**",
                            f"• {cta['cta_placement_strategy'][:200]}...",
                            ""
                        ])
                
                response_parts.extend([
                    f"💾 **Stored for Analysis** - Use `scrape insights` for competitive intelligence",
                    f"📈 **View History** - Use `scrape history` to see all analyzed sites"
                ])
                
                return "\n".join(response_parts)
                
            except Exception as e:
                logger.error(f"Scraper processing failed: {e}")
                return f"""❌ **Analysis Failed**
                
Error analyzing {url}: {str(e)}

Please try again or contact support if the issue persists."""
    
    except Exception as e:
        logger.error(f"Scraper command processing failed: {e}")
        return f"❌ **Scraper Command Error:** {str(e)}\n\nTry `scrape https://example.com` to analyze a website."

#-- Section 7: Prayer Times Functions - 9/26/25 (Updated 9/27/25)
def detect_prayer_command(message: str) -> bool:
    """Detect prayer-related requests"""
    prayer_keywords = [
        "prayer", "prayers", "salah", "namaz", "fajr", "dhuhr", "asr", "maghrib", "isha",
        "prayer time", "prayer times", "next prayer", "when is prayer", "how long until",
        "how long till", "time until prayer", "prayer schedule", "islamic time", "islamic date"
    ]
    message_lower = message.lower()
    return any(keyword in message_lower for keyword in prayer_keywords)

def detect_prayer_question_type(message: str) -> str:
    """Determine what type of prayer question the user is asking"""
    message_lower = message.lower()
    
    if any(phrase in message_lower for phrase in ["how long", "time until", "time till", "when is next"]):
        return "next_prayer"
    elif any(phrase in message_lower for phrase in ["prayer times today", "today's prayer", "prayer schedule", "all prayer"]):
        return "daily_schedule"
    elif any(phrase in message_lower for phrase in ["islamic date", "hijri", "islamic calendar"]):
        return "islamic_date"
    else:
        return "general_prayer"

async def process_prayer_command(message: str, user_id: str, ip_address: str = None) -> str:
    """Process prayer-related commands using the cached database system"""
    try:
        from ..integrations.prayer_times.database_manager import get_prayer_database_manager
        
        question_type = detect_prayer_question_type(message)
        
        # Get prayer manager and pass IP address for location detection
        prayer_manager = await get_prayer_database_manager()
        
        if question_type == "next_prayer":
            # "How long till Dhuhr?" type questions
            prayer_data = await prayer_manager.get_todays_prayer_times(ip_address)
            if not prayer_data:
                return """🕌 **Prayer Time Service Unavailable**
                
Unable to retrieve prayer times at the moment. Please try again in a few moments."""
            
            next_prayer_info = await prayer_manager.get_next_prayer_info()
            location_name = prayer_data['location']['name']
            
            if not next_prayer_info:
                return """🕌 **Prayer Time Service Unavailable**
                
Unable to retrieve prayer times at the moment. Please try again in a few moments."""
            
            prayer_name = next_prayer_info['prayer_name']
            prayer_time = next_prayer_info['prayer_time']
            time_until_text = next_prayer_info['time_until_text']
            is_today = next_prayer_info['is_today']
            
            day_text = "today" if is_today else "tomorrow"
            
            return f"""🕌 **Next Prayer: {prayer_name}**

⏰ **Time:** {prayer_time} ({day_text})
⏳ **Time Until:** {time_until_text}

Prayer times are calculated for {location_name} using ISNA method."""
        
        elif question_type == "daily_schedule":
            # "What are prayer times today?" type questions
            prayer_data = await prayer_manager.get_todays_prayer_times(ip_address)
            
            if not prayer_data:
                return """🕌 **Prayer Schedule Unavailable**
                
Unable to retrieve today's prayer schedule. Please try again in a few moments."""
            
            prayer_times = prayer_data['prayer_times']
            islamic_date = prayer_data['islamic_date']
            formatted_date = prayer_data.get('formatted_date', prayer_data['date'])
            location_name = prayer_data['location']['name']
            
            response_parts = [
                f"🕌 **Prayer Times for {formatted_date}**",
                ""
            ]
            
            # Add Islamic date if available
            if islamic_date['date'] != 'N/A':
                response_parts.extend([
                    f"📅 **Islamic Date:** {islamic_date['date']} {islamic_date['month']} {islamic_date['year']}",
                    ""
                ])
            
            response_parts.extend([
                "🕌 **Daily Prayer Schedule:**",
                f"   **Fajr:** {prayer_times['fajr']}",
                f"   **Dhuhr:** {prayer_times['dhuhr']}",
                f"   **Asr:** {prayer_times['asr']}",
                f"   **Maghrib:** {prayer_times['maghrib']}",
                f"   **Isha:** {prayer_times['isha']}",
                "",
                f"📍 Calculated for {location_name} using ISNA method"
            ])
            
            return "\n".join(response_parts)
        
        elif question_type == "islamic_date":
            # Islamic calendar questions
            prayer_data = await prayer_manager.get_todays_prayer_times(ip_address)
            
            if not prayer_data or prayer_data['islamic_date']['date'] == 'N/A':
                return """📅 **Islamic Calendar Information**
                
Islamic date information is currently unavailable. Please try again later."""
            
            islamic_date = prayer_data['islamic_date']
            gregorian_date = prayer_data.get('formatted_date', prayer_data['date'])
            
            return f"""📅 **Islamic Calendar Information**

**Today's Date:**
📆 Gregorian: {gregorian_date}
🗓️ Islamic: {islamic_date['date']} {islamic_date['month']} {islamic_date['year']}

Islamic dates are calculated using the AlAdhan calendar system."""
        
        else:
            # General prayer information
            prayer_data = await prayer_manager.get_todays_prayer_times(ip_address)
            next_prayer_info = await prayer_manager.get_next_prayer_info()
            
            if not next_prayer_info or not prayer_data:
                return """🕌 **Prayer Time Information**
                
Prayer time service is currently unavailable. Please try again in a few moments."""
            
            next_prayer = next_prayer_info['prayer_name']
            time_until = next_prayer_info['time_until_text']
            islamic_date = prayer_data['islamic_date']
            location_name = prayer_data['location']['name']
            
            response_parts = [
                "🕌 **Prayer Time Information**",
                "",
                f"⏰ **Next Prayer:** {next_prayer} in {time_until}",
                f"📍 **Location:** {location_name}",
                ""
            ]
            
            if islamic_date['date'] != 'N/A':
                response_parts.extend([
                    f"📅 **Islamic Date:** {islamic_date['date']} {islamic_date['month']} {islamic_date['year']}",
                    ""
                ])
            
            response_parts.extend([
                "**Available Commands:**",
                "• 'How long till [prayer name]?' - Next prayer countdown",
                "• 'What are prayer times today?' - Full daily schedule",
                "• 'Islamic date' - Current Hijri calendar date"
            ])
            
            return "\n".join(response_parts)
    
    except Exception as e:
        logger.error(f"Prayer command processing failed: {e}")
        return f"""🕌 **Prayer Time Service Error**
        
An error occurred while retrieving prayer information: {str(e)}

Please try again or contact support if the issue persists."""

#-- Section 8: File Processing Functions - 9/26/25
async def process_uploaded_files(files) -> List[Dict]:
    """Process uploaded files and return file information"""
    from fastapi import UploadFile
    
    processed_files = []
    
    # Ensure upload directory exists
    ensure_upload_dir()
    
    for file in files:
        if not file.filename:
            continue
            
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in ALLOWED_EXTENSIONS:
            raise Exception(f"File type {file_ext} not allowed. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}")
        
        content = await file.read()
        if len(content) > MAX_FILE_SIZE:
            raise Exception(f"File {file.filename} too large. Max size: 10MB")
        
        file_id = str(uuid.uuid4())
        file_path = UPLOAD_DIR / f"{file_id}_{file.filename}"
        
        with open(file_path, "wb") as f:
            f.write(content)
        
        file_info = {
            'file_id': file_id,
            'filename': file.filename,
            'file_type': file_ext,
            'file_size': len(content),
            'file_path': str(file_path),
            'analysis': await analyze_file_content(file_path, file_ext)
        }
        
        processed_files.append(file_info)
        await file.seek(0)
    
    return processed_files

async def analyze_file_content(file_path: Path, file_type: str) -> Dict:
    """Analyze file content and extract information"""
    analysis = {
        'type': 'unknown',
        'description': '',
        'extracted_text': '',
        'metadata': {},
        'key_insights': []
    }
    
    try:
        if file_type in ['.png', '.jpg', '.jpeg', '.gif']:
            analysis.update(await analyze_image_file(file_path))
        elif file_type == '.pdf':
            analysis.update(await analyze_pdf_file(file_path))
        elif file_type in ['.txt', '.md']:
            analysis.update(await analyze_text_file(file_path))
        elif file_type == '.csv':
            analysis.update(await analyze_csv_file(file_path))
        elif file_type in ['.doc', '.docx']:
            analysis.update(await analyze_docx_file(file_path))
        elif file_type in ['.xls', '.xlsx']:
            analysis.update(await analyze_excel_file(file_path))
        elif file_type == '.py':
            analysis.update(await analyze_python_file(file_path))
        else:
            analysis['description'] = f"Unsupported file type: {file_type}"
            
    except Exception as e:
        logger.error(f"File analysis failed for {file_path}: {e}")
        analysis['description'] = f"Analysis failed: {str(e)}"
        analysis['error'] = str(e)
    
    return analysis

async def analyze_image_file(file_path: Path) -> Dict:
    """Analyze image file"""
    try:
        with Image.open(file_path) as img:
            width, height = img.size
            format_name = img.format
            mode = img.mode
            
        return {
            'type': 'image',
            'description': f'{format_name} image ({width}x{height}, {mode})',
            'metadata': {
                'width': width,
                'height': height,
                'format': format_name,
                'mode': mode
            },
            'extracted_text': f"Image file: {width}x{height} {format_name}"
        }
    except Exception as e:
        return {
            'type': 'image',
            'description': f'Image analysis failed: {e}',
            'extracted_text': ''
        }

async def analyze_pdf_file(file_path: Path) -> Dict:
    """Analyze PDF file"""
    try:
        text_content = ""
        page_count = 0
        
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages[:5]:  # First 5 pages
                if page.extract_text():
                    text_content += page.extract_text() + "\n"
        
        return {
            'type': 'pdf',
            'description': f'PDF document with {page_count} pages',
            'extracted_text': text_content[:2000],  # First 2000 chars
            'metadata': {'page_count': page_count}
        }
    except Exception as e:
        return {
            'type': 'pdf',
            'description': f'PDF analysis failed: {e}',
            'extracted_text': ''
        }

async def analyze_text_file(file_path: Path) -> Dict:
    """Analyze text file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        line_count = len(content.splitlines())
        word_count = len(content.split())
        
        return {
            'type': 'text',
            'description': f'Text file ({line_count} lines, {word_count} words)',
            'extracted_text': content[:2000],  # First 2000 chars
            'metadata': {'line_count': line_count, 'word_count': word_count}
        }
    except Exception as e:
        return {
            'type': 'text',
            'description': f'Text analysis failed: {e}',
            'extracted_text': ''
        }

async def analyze_csv_file(file_path: Path) -> Dict:
    """Analyze CSV file"""
    try:
        df = pd.read_csv(file_path)
        rows, cols = df.shape
        columns = df.columns.tolist()
        
        preview = df.head().to_string()
        
        return {
            'type': 'csv',
            'description': f'CSV file ({rows} rows, {cols} columns)',
            'extracted_text': f"CSV Preview:\n{preview}",
            'metadata': {'rows': rows, 'columns': cols, 'column_names': columns}
        }
    except Exception as e:
        return {
            'type': 'csv',
            'description': f'CSV analysis failed: {e}',
            'extracted_text': ''
        }
        
async def analyze_docx_file(file_path: Path) -> Dict:
    """Analyze Word document file (.docx)"""
    try:
        doc = Document(file_path)
        
        # Extract text content
        text_content = ""
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        # Count paragraphs, tables, images
        paragraph_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        
        # Extract table data
        table_text = ""
        for table in doc.tables[:3]:  # First 3 tables
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " | "
                table_text += "\n"
        
        word_count = len(text_content.split())
        
        return {
            'type': 'document',
            'description': f'Word document with {paragraph_count} paragraphs, {table_count} tables, {word_count} words',
            'extracted_text': (text_content + "\n\n" + table_text)[:3000],  # First 3000 chars
            'metadata': {
                'paragraph_count': paragraph_count,
                'table_count': table_count,
                'word_count': word_count
            }
        }
    except Exception as e:
        logger.error(f"Word document analysis failed: {e}")
        return {
            'type': 'document',
            'description': f'Word document analysis failed: {e}',
            'extracted_text': ''
        }
        
async def analyze_excel_file(file_path: Path) -> Dict:
    """Analyze Excel file (.xlsx, .xls)"""
    try:
        # Load workbook
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        
        sheet_names = workbook.sheetnames
        sheet_count = len(sheet_names)
        
        # Analyze first sheet
        first_sheet = workbook.active
        row_count = first_sheet.max_row
        col_count = first_sheet.max_column
        
        # Extract sample data (first 10 rows)
        extracted_text = f"Sheet: {first_sheet.title}\n\n"
        for row_idx, row in enumerate(first_sheet.iter_rows(max_row=10, values_only=True), 1):
            row_data = [str(cell) if cell is not None else '' for cell in row]
            extracted_text += " | ".join(row_data) + "\n"
            if row_idx >= 10:
                break
        
        # Get all sheet summaries
        sheet_summaries = []
        for sheet_name in sheet_names[:5]:  # First 5 sheets
            sheet = workbook[sheet_name]
            sheet_summaries.append(f"{sheet_name}: {sheet.max_row} rows × {sheet.max_column} cols")
        
        return {
            'type': 'spreadsheet',
            'description': f'Excel file with {sheet_count} sheets, {row_count} rows × {col_count} columns in active sheet',
            'extracted_text': extracted_text[:3000],
            'metadata': {
                'sheet_count': sheet_count,
                'sheet_names': sheet_names[:10],  # First 10 sheet names
                'active_sheet': first_sheet.title,
                'row_count': row_count,
                'column_count': col_count,
                'sheet_summaries': sheet_summaries
            }
        }
    except Exception as e:
        logger.error(f"Excel file analysis failed: {e}")
        return {
            'type': 'spreadsheet',
            'description': f'Excel analysis failed: {e}',
            'extracted_text': ''
        }
        
async def analyze_python_file(file_path: Path) -> Dict:
    """Analyze Python source code file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            source_code = f.read()
        
        # Parse AST for code analysis
        try:
            tree = ast.parse(source_code)
            
            # Count different code elements
            functions = [node.name for node in ast.walk(tree) if isinstance(node, ast.FunctionDef)]
            classes = [node.name for node in ast.walk(tree) if isinstance(node, ast.ClassDef)]
            imports = []
            
            for node in ast.walk(tree):
                if isinstance(node, ast.Import):
                    imports.extend([alias.name for alias in node.names])
                elif isinstance(node, ast.ImportFrom):
                    if node.module:
                        imports.append(node.module)
            
            function_count = len(functions)
            class_count = len(classes)
            import_count = len(set(imports))
            
        except SyntaxError:
            function_count = class_count = import_count = 0
            functions = classes = imports = []
        
        # Count lines
        lines = source_code.splitlines()
        line_count = len(lines)
        code_lines = [line for line in lines if line.strip() and not line.strip().startswith('#')]
        comment_lines = [line for line in lines if line.strip().startswith('#')]
        
        # Extract docstrings and top comments
        extracted_text = source_code[:2000]  # First 2000 chars
        
        return {
            'type': 'code',
            'description': f'Python file with {function_count} functions, {class_count} classes, {line_count} lines',
            'extracted_text': extracted_text,
            'metadata': {
                'language': 'python',
                'line_count': line_count,
                'code_lines': len(code_lines),
                'comment_lines': len(comment_lines),
                'function_count': function_count,
                'class_count': class_count,
                'import_count': import_count,
                'functions': functions[:10],  # First 10 function names
                'classes': classes[:10],  # First 10 class names
                'imports': list(set(imports))[:15]  # First 15 unique imports
            }
        }
    except Exception as e:
        logger.error(f"Python file analysis failed: {e}")
        return {
            'type': 'code',
            'description': f'Python file analysis failed: {e}',
            'extracted_text': ''
        }

#-- Section 9: Voice Synthesis Functions - 9/28/25
def detect_voice_command(message: str) -> bool:
    """Detect voice synthesis commands"""
    voice_keywords = [
        "voice synthesize", "voice generate", "say this", "speak this",
        "voice this", "read this", "voice history", "voice personalities",
        "text to speech", "tts", "synthesize voice", "generate audio"
    ]
    message_lower = message.lower()
    return any(keyword in message_lower for keyword in voice_keywords)

def extract_text_for_voice(message: str) -> str:
    """Extract text to synthesize from voice command"""
    message_lower = message.lower()
    
    # Common patterns for voice commands
    patterns = [
        r'voice synthesize (.+)',
        r'voice generate (.+)',
        r'say this[:\s]+(.+)',
        r'speak this[:\s]+(.+)',
        r'voice this[:\s]+(.+)',
        r'read this[:\s]+(.+)',
        r'synthesize (.+)',
        r'tts (.+)'
    ]
    
    import re
    for pattern in patterns:
        match = re.search(pattern, message_lower)
        if match:
            return match.group(1).strip()
    
    # If no pattern matches, check if message is short enough to be direct synthesis
    if len(message.split()) <= 50:  # Reasonable voice synthesis length
        return message
    
    return None

async def process_voice_command(message: str, user_id: str) -> str:
    """Process voice synthesis commands"""
    try:
        from ..integrations.voice_synthesis import (
            get_voice_client,
            get_personality_voice_manager,
            get_audio_cache_manager
        )
        
        message_lower = message.lower()
        
        if 'voice history' in message_lower:
            # Get voice synthesis history
            audio_manager = get_audio_cache_manager()
            stats = await audio_manager.get_cache_statistics()
            
            recent_generations = stats.get('recent_generations_24h', 0)
            total_files = stats.get('total_files', 0)
            total_size_mb = stats.get('total_size_mb', 0.0)
            
            return f"""🎤 **Voice Synthesis History**

📊 **Recent Activity:** {recent_generations} audio files generated today
💾 **Total Generated:** {total_files} audio files
📈 **Storage Used:** {total_size_mb:.1f} MB

**Available Commands:**
• `voice synthesize [text]` - Generate speech from text
• `voice personalities` - View available voices
• Click speaker buttons on AI responses for instant audio

Ready to give your words a voice?"""
        
        elif 'voice personalities' in message_lower or 'voice list' in message_lower:
            # Show personality voice mappings
            personality_manager = get_personality_voice_manager()
            mappings = await personality_manager.get_all_personality_mappings()
            
            response_parts = ["🎭 **Personality Voice Mappings**\n"]
            
            for personality_id, mapping in mappings.items():
                primary_voice = mapping.get('primary_voice', 'Unknown')
                style = mapping.get('voice_settings', {}).get('style', 'Standard')
                
                response_parts.append(f"**{personality_id.title()}:** {primary_voice}")
                response_parts.append(f"   Style: {style.replace('_', ' ').title()}")
                response_parts.append("")
            
            response_parts.append("Each personality has been carefully matched with voices chosen by SyntaxPrime for optimal expression!")
            return "\n".join(response_parts)
        
        else:
            # Extract text to synthesize
            text_to_synthesize = extract_text_for_voice(message)
            
            if not text_to_synthesize:
                return """🎤 **Voice Synthesis Commands**

**Usage:**
• `voice synthesize [your text here]` - Generate speech from text
• `voice history` - View synthesis history
• `voice personalities` - See available voices

**Examples:**
• `voice synthesize Hello everyone, welcome to the presentation`
• `say this: Thanks for joining our meeting today`
• `read this: The quarterly results show significant growth`

**Features:**
✅ Personality-specific voices (SyntaxPrime's choices)
✅ High-quality MP3 generation
✅ Instant inline playback
✅ Database caching for re-use

Ready to bring your words to life?"""
            
            # Synthesize the audio
            voice_client = get_voice_client()
            personality_manager = get_personality_voice_manager()
            audio_manager = get_audio_cache_manager()
            
            # Get voice for current context (default to syntaxprime)
            personality_id = 'syntaxprime'  # Could be determined from context
            voice_id = personality_manager.get_voice_for_personality(personality_id)
            
            # Generate speech
            synthesis_result = await voice_client.generate_speech(
                text=text_to_synthesize,
                voice_id=voice_id,
                personality_id=personality_id
            )
            
            if not synthesis_result.get('success'):
                return f"""❌ **Voice Synthesis Failed**

Unable to generate audio for: "{text_to_synthesize[:50]}..."

Error: {synthesis_result.get('error', 'Unknown error')}

Please try again or contact support if the issue persists."""
            
            # Create a message ID for caching
            message_id = str(uuid.uuid4())
            
            # Cache the audio
            cache_result = await audio_manager.cache_audio(
                message_id=message_id,
                audio_data=synthesis_result['audio_data'],
                metadata={
                    'personality_id': personality_id,
                    'voice_id': voice_id,
                    'text_length': len(text_to_synthesize),
                    'generation_time_ms': synthesis_result.get('generation_time_ms'),
                    'file_format': 'mp3'
                }
            )
            
            generation_time = synthesis_result.get('generation_time_seconds', 0)
            file_size = len(synthesis_result['audio_data'])
            voice_name = personality_manager._get_voice_name_from_id(voice_id)
            
            return f"""✅ **Voice Synthesis Complete**

🎤 **Text:** "{text_to_synthesize[:100]}{'...' if len(text_to_synthesize) > 100 else ''}"
🗣️ **Voice:** {voice_name} ({personality_id} personality)
⏱️ **Generation Time:** {generation_time:.1f} seconds
📁 **File Size:** {file_size:,} bytes

🔊 **Audio URL:** `/api/voice/audio/{message_id}`

The audio has been generated and cached for playback. Use the audio URL in your frontend to play the synthesized speech!"""
    
    except Exception as e:
        logger.error(f"Voice command processing failed: {e}")
        return f"❌ **Voice Synthesis Error:** {str(e)}\n\nTry `voice synthesize hello world` to test the system."

#-- Section 10: Image Generation Functions - 9/28/25
def detect_image_command(message: str) -> bool:
    """Detect image generation commands"""
    image_keywords = [
        "image create", "image generate", "generate image", "create image",
        "image blog", "image social", "image marketing", "image history",
        "image download", "make image", "draw image", "picture of",
        "visualize this", "show me", "image style", "mockup"
    ]
    message_lower = message.lower()
    return any(keyword in message_lower for keyword in image_keywords)

def extract_image_prompt(message: str) -> tuple[str, str]:
    """Extract image prompt and content type from command"""
    message_lower = message.lower()
    
    # Detect content type from command
    content_type = 'general'
    if 'blog' in message_lower:
        content_type = 'blog'
    elif 'social' in message_lower:
        content_type = 'social'
    elif 'marketing' in message_lower:
        content_type = 'marketing'
    elif 'illustration' in message_lower:
        content_type = 'illustration'
    
    # Extract prompt text
    import re
    patterns = [
        r'image create (.+)',
        r'image generate (.+)',
        r'generate image (.+)',
        r'create image (.+)',
        r'image blog (.+)',
        r'image social (.+)',
        r'image marketing (.+)',
        r'make image (.+)',
        r'draw (.+)',
        r'picture of (.+)',
        r'visualize (.+)',
        r'show me (.+)'
        r'mockup (.+)'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, message_lower)
        if match:
            return match.group(1).strip(), content_type
    
    # If no pattern matches but it's an image command, use whole message as prompt
    if detect_image_command(message):
        return message.strip(), content_type
    
    return None, content_type

async def process_image_command(message: str, user_id: str) -> str:
    """Process image generation commands"""
    try:
        from ..integrations.image_generation import (
            generate_image_for_chat,
            get_image_history_for_chat
        )
        from ..integrations.image_generation.database_manager import ImageDatabase
        
        message_lower = message.lower()
        
        if 'image history' in message_lower:
            # Get image generation history
            history = await get_image_history_for_chat(user_id, limit=10)
            
            if not history:
                return """🎨 **Image Generation History**

No images generated yet. Start creating with:
• `image create [description]` - Generate any image
• `image blog [topic]` - Create blog header image
• `image social [content]` - Generate social media graphic

Ready to bring your ideas to life?"""
            
            response_parts = ["🖼️ **Recent Generated Images**\n"]
            
            for i, img in enumerate(history, 1):
                response_parts.append(f"**{i}. {img['content_type'].title()} Image**")
                response_parts.append(f"   📝 Prompt: {img['prompt']}")
                response_parts.append(f"   🗓️ Created: {img['created_at']}")
                response_parts.append(f"   📥 Downloads: {img['download_count']}")
                response_parts.append("")
            
            response_parts.append("💡 Use `image create [new prompt]` to generate more images!")
            return "\n".join(response_parts)
        
        else:
            # Extract prompt and generate image
            prompt, content_type = extract_image_prompt(message)
            
            if not prompt:
                return """🎨 **Image Generation Commands**

**Usage:**
• `image create [description]` - Generate image from description
• `image blog [topic]` - Create blog featured image
• `image social [content]` - Generate social media graphic
• `image marketing [campaign]` - Create marketing visual
• `image history` - View generated images

**Examples:**
• `image create a professional business meeting in a modern office`
• `image blog content marketing trends for 2024`
• `image social new product announcement celebration`
• `image marketing email campaign header design`

**Features:**
✅ Inline base64 display (no external URLs)
✅ Smart model selection for quality
✅ Multiple download formats
✅ Content-type optimization
✅ Style templates available

Ready to create stunning visuals?"""
            
            # Generate the image
            generation_result = await generate_image_for_chat(
                prompt=prompt,
                content_type=content_type,
                user_id=user_id
            )
            
            if not generation_result.get('success'):
                return f"""❌ **Image Generation Failed**

Unable to create image for: "{prompt[:50]}..."

Error: {generation_result.get('error', 'Unknown error')}

Please try a different prompt or contact support if the issue persists."""
            
            # Successful generation
            image_id = generation_result.get('image_id', 'unknown')
            model_used = generation_result.get('model_used', 'Unknown')
            generation_time = generation_result.get('generation_time_seconds', 0)
            enhanced_prompt = generation_result.get('enhanced_prompt', prompt)
            resolution = generation_result.get('resolution', '1024x1024')
            image_base64 = generation_result.get('image_base64', '')
            
            return f"""✅ **Image Generated Successfully**

🎨 **Original Prompt:** "{prompt}"
🧠 **Enhanced Prompt:** "{enhanced_prompt[:100]}{'...' if len(enhanced_prompt) > 100 else ''}"
🤖 **Model Used:** {model_used}
📐 **Resolution:** {resolution}
⏱️ **Generation Time:** {generation_time:.1f} seconds
💾 **Image ID:** {image_id}

🖼️ **Image displayed below**

<IMAGE_DATA>{image_base64}</IMAGE_DATA>

**Download Options:**
📥 PNG: `/integrations/image-generation/download/{image_id}?format=png`
📥 JPG: `/integrations/image-generation/download/{image_id}?format=jpg`
📥 WebP: `/integrations/image-generation/download/{image_id}?format=webp`

Your visual idea has been brought to life! 🌟"""
    
    except Exception as e:
        logger.error(f"Image command processing failed: {e}")
        return f"❌ **Image Generation Error:** {str(e)}\n\nTry `image create blue circle` to test the system."

#-- Section 11: Google Trends Integration Functions - 9/27/25
def detect_trends_command(message: str) -> tuple[bool, str]:
    """Detect Google Trends commands and determine command type"""
    trends_keywords = [
        'trends', 'trending', 'google trends', 'opportunities',
        'good match', 'bad match', 'train trends', 'trends status',
        'trends health', 'trend opportunities', 'trends scan'
    ]
    
    message_lower = message.lower()
    
    # Check if it's a trends-related command
    is_trends_command = any(keyword in message_lower for keyword in trends_keywords)
    
    if not is_trends_command:
        return False, ''
    
    # Determine command type
    if any(term in message_lower for term in ['good match', 'bad match']):
        return True, 'training_feedback'
    elif 'opportunities' in message_lower:
        return True, 'view_opportunities'
    elif 'scan' in message_lower:
        return True, 'scan_trends'
    elif 'status' in message_lower or 'health' in message_lower:
        return True, 'status_check'
    else:
        return True, 'general_trends'

async def process_training_feedback(message: str, user_id: str) -> str:
    """Process Good Match/Bad Match training feedback"""
    try:
        training = OpportunityTraining()
        message_lower = message.lower()
        
        if 'good match' in message_lower:
            # User is providing positive feedback
            feedback_result = await training.record_feedback(
                user_id=user_id,
                feedback_type='positive',
                message_context=message
            )
            
            return f"""✅ **Training Feedback Recorded**

💎 **Feedback Type:** Good Match
📈 **Impact:** This helps improve trend opportunity detection
🤖 **ML Effect:** Future similar opportunities will be prioritized

**Your Training Stats:**
• Total Feedback: {feedback_result.get('total_feedback', 0)}
• Good Matches: {feedback_result.get('positive_feedback', 0)}
• Bad Matches: {feedback_result.get('negative_feedback', 0)}

Keep the feedback coming to improve accuracy!"""
        
        elif 'bad match' in message_lower:
            # User is providing negative feedback
            feedback_result = await training.record_feedback(
                user_id=user_id,
                feedback_type='negative',
                message_context=message
            )
            
            return f"""❌ **Training Feedback Recorded**

💎 **Feedback Type:** Bad Match
📉 **Impact:** This helps filter out irrelevant opportunities
🤖 **ML Effect:** Future similar opportunities will be deprioritized

**Your Training Stats:**
• Total Feedback: {feedback_result.get('total_feedback', 0)}
• Good Matches: {feedback_result.get('positive_feedback', 0)}
• Bad Matches: {feedback_result.get('negative_feedback', 0)}

Your feedback helps make the system smarter!"""
        
        else:
            return """🤖 **Training Feedback System**

Help improve trend opportunity detection with feedback:

**Available Commands:**
• `Good Match` - Mark current opportunity as relevant
• `Bad Match` - Mark current opportunity as irrelevant
• `trends opportunities` - View current opportunities to evaluate

**How Training Works:**
📈 Good Match feedback increases similar opportunity scores
📉 Bad Match feedback decreases similar opportunity scores
💎 More feedback = Better opportunity detection accuracy"""
    
    except Exception as e:
        logger.error(f"Training feedback processing failed: {e}")
        return f"❌ **Training Error:** {str(e)}\n\nTry `trends status` to check system health."

async def check_for_trend_opportunities(user_id: str) -> Optional[str]:
    """Check for proactive trend opportunities and return notification message"""
    try:
        detector = OpportunityDetector(database_url)
        
        # Get current trending opportunities
        opportunities = await detector.detect_current_opportunities(hours_lookback=24)
        
        if not opportunities:
            return None
        
        # Filter for high-confidence opportunities only
        high_confidence_ops = [
            op for op in opportunities
            if op.get('confidence_score', 0) >= 0.75
        ]
        
        if not high_confidence_ops:
            return None
        
        # Create proactive notification
        top_opportunity = high_confidence_ops[0]
        
        notification_message = f"""📈 **Trending Opportunity Alert**

**{top_opportunity.get('keyword', 'Unknown Trend')}** is gaining momentum!

📊 **Trend Score:** {top_opportunity.get('trend_score', 0):.1f}/10
💎 **Relevance:** {top_opportunity.get('confidence_score', 0):.0%}
📅 **Peak Expected:** {top_opportunity.get('peak_timing', 'Soon')}

**Why This Matters:**
{top_opportunity.get('opportunity_reason', 'High engagement potential detected')}

**Suggested Actions:**
• {top_opportunity.get('suggested_action_1', 'Create content around this trend')}
• {top_opportunity.get('suggested_action_2', 'Monitor for additional opportunities')}

Use `Good Match` or `Bad Match` to train the system!"""
        
        return notification_message
        
    except Exception as e:
        logger.error(f"Trend opportunity check failed: {e}")
        return None

async def process_trends_command(message: str, user_id: str) -> str:
    """Process Google Trends commands"""
    try:
        is_trends_cmd, cmd_type = detect_trends_command(message)
        
        if not is_trends_cmd:
            return ""
        
        if cmd_type == 'training_feedback':
            return await process_training_feedback(message, user_id)
        
        elif cmd_type == 'view_opportunities':
            detector = OpportunityDetector(database_url)
            opportunities = await detector.detect_current_opportunities(hours_lookback=24)
            
            if not opportunities:
                return """📈 **Google Trends Opportunities**

No trending opportunities found at this time.

**Available Commands:**
• `trends scan` - Force scan for new opportunities
• `trends status` - Check system health
• Use `Good Match`/`Bad Match` to train the system"""
            
            response_parts = ["📈 **Current Trending Opportunities**\n"]
            
            for i, opp in enumerate(opportunities, 1):
                response_parts.append(f"""**{i}. {opp.keyword}**
            📊 Trend Score: {opp.trend_score}/10
            💎 Relevance: {opp.opportunity_score:.0%}
            💡 Why: {opp.reasoning}

            """)
            
            response_parts.append("Respond with `Good Match` or `Bad Match` to help train the system!")
            return "\n".join(response_parts)
        
        elif cmd_type == 'scan_trends':
            # Force a new scan
            detector = OpportunityDetector(database_url)
            scan_result = await detector.force_scan_update()
            
            return f"""🔍 **Trends Scan Complete**

✅ Scanned {scan_result.get('trends_analyzed', 0)} trending topics
📈 Found {scan_result.get('opportunities_detected', 0)} new opportunities
⏰ Scan completed at {datetime.now().strftime('%H:%M')}

Use `trends opportunities` to view the latest findings!"""
        
        elif cmd_type == 'status_check':
            health_status = check_module_health()
            
            return f"""📈 **Google Trends System Status**

📡 **Service:** {'Running' if health_status.get('trends_healthy', False) else 'Issues Detected'}
⏰ **Last Scan:** {health_status.get('last_trends_scan', 'Unknown')}
📊 **Opportunities Found:** {health_status.get('total_opportunities', 0)}
🤖 **Training Data:** {health_status.get('training_samples', 0)} feedback samples

**Available Commands:**
• `trends opportunities` - View current opportunities
• `trends scan` - Force new scan
• `Good Match`/`Bad Match` - Provide training feedback"""
        
        else:
            # General trends information
            return """📈 **Google Trends Intelligence System**

🤖 **AI-Powered Opportunity Detection**
💎 **Personalized Content Suggestions**
📊 **Real-Time Trend Analysis**

**Available Commands:**
• `trends opportunities` - View current trending opportunities
• `trends scan` - Force scan for new trends
• `trends status` - Check system health
• `Good Match` - Mark opportunity as relevant (training)
• `Bad Match` - Mark opportunity as irrelevant (training)

**How It Works:**
1. Continuously monitors Google Trends
2. AI analyzes relevance to your interests
3. Presents high-potential opportunities
4. Learns from your feedback to improve accuracy

Ready to discover your next trending opportunity?"""
    
    except Exception as e:
        logger.error(f"Trends command processing failed: {e}")
        return f"❌ **Trends System Error:** {str(e)}\n\nTry `trends status` to check system health."

#-- Section 12: Prayer Notification Functions - 9/27/25
def detect_prayer_notification_command(message: str) -> bool:
    """Detect prayer notification management commands"""
    notification_keywords = [
        "prayer notifications", "prayer alerts", "prayer reminder",
        "notification status", "notification test", "prayer service",
        "disable prayer", "enable prayer", "prayer settings"
    ]
    message_lower = message.lower()
    return any(keyword in message_lower for keyword in notification_keywords)

async def process_prayer_notification_command(message: str, user_id: str, ip_address: str = None) -> str:
    """Process prayer notification management commands"""
    try:
        from ..integrations.prayer_times.notification_manager import (
            get_prayer_notification_manager,
            test_prayer_notification
        )
        
        message_lower = message.lower()
        notification_manager = get_prayer_notification_manager()
        
        if 'status' in message_lower:
            # Get notification service status
            status = notification_manager.get_notification_status()
            
            return f"""🕌 **Prayer Notification Service Status**

📡 **Service:** {'Running' if status['running'] else 'Stopped'}
⏰ **Check Interval:** {status['check_interval_seconds']} seconds
📅 **Advance Notice:** {status['advance_minutes']} minutes
📊 **Notifications Sent Today:** {status['sent_today']}

**Preferences:**
- Enabled: {'Yes' if status['preferences']['enabled'] else 'No'}
- Advance Time: {status['preferences']['advance_minutes']} minutes
- Personality: {'Enabled' if status['preferences']['personality_enabled'] else 'Disabled'}
- Prayers: {', '.join(status['preferences']['prayers_to_notify']).title()}

Use `prayer notifications test` to test the system."""
        
        elif 'test' in message_lower:
            # Test the notification system
            test_message = await test_prayer_notification()
            
            return f"""🧪 **Prayer Notification Test**

Here's what a notification would look like:

---

{test_message}

---

If the service is running properly, you should receive automatic notifications 15 minutes before each prayer time."""
        
        else:
            # General prayer notification info
            return """🕌 **Prayer Notification System**

**Available Commands:**
- `prayer notifications status` - Check service status
- `prayer notifications test` - Test notification format
- `prayer times` - Get today's prayer schedule
- `next prayer` - See time until next prayer

**Features:**
✅ Automatic 15-minute advance notifications
✅ AI personality integration for natural messages
✅ All 5 daily prayers included
✅ No duplicate notifications

The service runs automatically in the background."""
            
    except Exception as e:
        logger.error(f"Prayer notification command processing failed: {e}")
        return f"❌ **Prayer Notification Error:** {str(e)}\n\nTry `prayer notifications status` to check the system."

def detect_location_command(message: str) -> bool:
    """Detect location-related commands"""
    location_keywords = [
        "my location", "where am i", "current location", "detect location",
        "prayer location", "location for prayers", "change location",
        "location settings", "ip location", "auto location"
    ]
    message_lower = message.lower()
    return any(keyword in message_lower for keyword in location_keywords)

async def process_location_command(message: str, user_id: str, ip_address: str = None) -> str:
    """Process location detection and management commands"""
    try:
        from ..integrations.prayer_times.location_detector import (
            get_location_detector,
            detect_user_location
        )
        
        message_lower = message.lower()
        detector = get_location_detector()
        
        if any(term in message_lower for term in ['where am i', 'current location', 'my location']):
            # Detect and show current location
            location = await detect_user_location(ip_address)
            
            return f"""📍 **Your Current Location**

🌐 **Detected from IP:** {location.get('source', 'IP Service')}
🏙️ **City:** {location['city']}
📍 **Region:** {location['region']}
🌍 **Country:** {location['country']}
📊 **Coordinates:** {location['latitude']:.4f}, {location['longitude']:.4f}
🕐 **Timezone:** {location['timezone']}

This location will be used automatically for prayer time calculations."""
        
        elif 'prayer location' in message_lower:
            # Get location specifically formatted for prayers
            from ..integrations.prayer_times.location_detector import get_prayer_location
            location_name, lat, lng = await get_prayer_location(user_id, ip_address)
            
            return f"""🕌 **Prayer Times Location**

📍 **Location:** {location_name}
📊 **Coordinates:** {lat:.4f}, {lng:.4f}

Prayer times are automatically calculated for your current location based on your IP address. 

**Available Commands:**
- `prayer times` - Get today's schedule for your location
- `next prayer` - Time until next prayer
- `my location` - See detected location details"""
        
        else:
            # General location information
            return f"""📍 **Location Detection System**

Your prayer times are automatically calculated based on your IP address location.

**Available Commands:**
- `my location` - See your detected location
- `prayer location` - View location used for prayers  
- `location settings` - Check auto-detection settings
- `prayer times` - Get schedule for your current location

**Features:**
✅ Automatic IP-based location detection
✅ 24-hour location caching for performance
✅ Multiple geolocation services for accuracy
✅ Fallback to default location if needed"""
            
    except Exception as e:
        logger.error(f"Location command processing failed: {e}")
        return f"❌ **Location Detection Error:** {str(e)}\n\nTry `my location` to test location detection."

async def post_system_message_to_chat(message: str, message_type: str = "system_notification") -> bool:
    """
    Post a system-generated message to the chat interface
    This is used by background services like prayer notifications
    """
    try:
        from .conversation_manager import get_memory_manager
        
        # Get the default user ID and create a system thread if needed
        DEFAULT_USER_ID = "b7c60682-4815-4d9d-8ebe-66c6cd24eff9"
        
        # Get memory manager
        memory_manager = get_memory_manager(DEFAULT_USER_ID)
        
        # Create or get system notification thread
        system_thread_id = await get_or_create_system_thread(memory_manager)
        
        # Add system message to conversation history
        system_message_id = await memory_manager.add_message(
            thread_id=system_thread_id,
            role="assistant",  # System messages appear as assistant responses
            content=message,
            metadata={
                "type": message_type,
                "timestamp": datetime.now().isoformat(),
                "source": "prayer_notification_service"
            }
        )
        
        logger.info(f"✅ System message posted to chat: {system_message_id}")
        return True
        
    except Exception as e:
        logger.error(f"Failed to post system message to chat: {e}")
        return False

async def get_or_create_system_thread(memory_manager) -> str:
    """Get or create a dedicated thread for system notifications"""
    try:
        # Create a new thread titled "Prayer Notifications"
        thread_id = await memory_manager.create_conversation_thread(
            platform='system',
            title='Prayer Notifications'
        )
        
        return thread_id
        
    except Exception as e:
        logger.error(f"Failed to create system thread: {e}")
        raise

#-- Section 13: Module Information Functions - 9/26/25 (Updated 9/28/25)
def get_integration_info():
    """Get information about the chat integration helper module"""
    return {
        "name": "AI Chat Integration Helper",
        "version": "2.3.0",
        "description": "Helper functions for weather, prayer times, scraper, RSS, Bluesky, Google Trends, Voice Synthesis, Image Generation, and file processing",
        "note": "This module provides helper functions only - endpoints are handled by router.py",
        "features": [
            "Weather integration with Tomorrow.io API",
            "Prayer times with AlAdhan API integration and notifications",
            "Marketing scraper for competitive analysis",
            "RSS learning integration for marketing insights",
            "Bluesky social media command processing",
            "Google Trends opportunity detection and training",
            "Voice synthesis with ElevenLabs integration",
            "Image generation with Replicate AI",
            "Google Workspace integration (Analytics, Search Console, Gmail, Calendar, Drive)",
            "Self-evolving Intelligence Engine with pattern recognition",
            "Multi-account Google authentication with OAuth device flow",
            "File upload processing (images, PDFs, CSVs, text)",
            "Real-time datetime context generation",
            "Location detection for prayer times",
            "Prayer notification management"
        ],
        "integrations_provided": [
            "weather_detection_and_processing",
            "prayer_times_detection_and_processing",
            "prayer_notification_management",
            "location_detection_and_processing",
            "marketing_scraper_detection_and_processing",
            "rss_learning_context_generation",
            "bluesky_command_processing",
            "google_trends_integration",
            "voice_synthesis_detection_and_processing",
            "image_generation_detection_and_processing",
            "google_workspace_command_processing",
            "file_upload_and_analysis"
        ]
    }

def check_module_health() -> Dict[str, Any]:
    """Check the health of the AI chat helper module"""
    missing_vars = []
    warnings = []
    
    if not UPLOAD_DIR.exists():
        warnings.append("Upload directory not found")
    
    if not os.getenv("TOMORROW_IO_API_KEY"):
        warnings.append("Weather integration not configured (TOMORROW_IO_API_KEY missing)")
    
    bluesky_configured = any([
        os.getenv("BLUESKY_PERSONAL_PASSWORD"),
        os.getenv("BLUESKY_ROSE_ANGEL_PASSWORD"),
        os.getenv("BLUESKY_BINGE_TV_PASSWORD"),
        os.getenv("BLUESKY_MEALS_FEELZ_PASSWORD"),
        os.getenv("BLUESKY_DAMN_IT_CARL_PASSWORD")
    ])
    
    if not bluesky_configured:
        warnings.append("Bluesky integration not configured (account passwords missing)")
    
    rss_configured = bool(os.getenv("DATABASE_URL"))
    
    if not rss_configured:
        warnings.append("RSS Learning integration not configured (DATABASE_URL missing)")
    
    # Check marketing scraper configuration
    scraper_configured = bool(os.getenv("DATABASE_URL"))  # Uses same DB as RSS
    
    if not scraper_configured:
        warnings.append("Marketing Scraper integration not configured (DATABASE_URL missing)")
    
    # Check prayer times configuration
    prayer_configured = bool(os.getenv("DATABASE_URL"))  # Uses same DB as others
    
    if not prayer_configured:
        warnings.append("Prayer Times integration not configured (DATABASE_URL missing)")
    
    # Check Google Trends configuration
    trends_configured = bool(os.getenv("DATABASE_URL"))  # Uses same DB as others
    
    if not trends_configured:
        warnings.append("Google Trends integration not configured (DATABASE_URL missing)")
    
    # Check Voice Synthesis configuration - NEW 9/28/25
    voice_configured = bool(os.getenv("ELEVENLABS_API_KEY"))
    
    if not voice_configured:
        warnings.append("Voice Synthesis integration not configured (ELEVENLABS_API_KEY missing)")
    
    # Check Image Generation configuration - NEW 9/28/25
    image_configured = bool(os.getenv("REPLICATE_API_TOKEN"))
    
    if not image_configured:
        warnings.append("Image Generation integration not configured (REPLICATE_API_TOKEN missing)")
    
    return {
        "healthy": len(missing_vars) == 0,
        "missing_vars": missing_vars,
        "warnings": warnings,
        "upload_directory": str(UPLOAD_DIR),
        "max_file_size": f"{MAX_FILE_SIZE // (1024 * 1024)}MB",
        "weather_integration_available": bool(os.getenv("TOMORROW_IO_API_KEY")),
        "bluesky_integration_available": bluesky_configured,
        "rss_learning_integration_available": rss_configured,
        "marketing_scraper_integration_available": scraper_configured,
        "prayer_times_integration_available": prayer_configured,
        "google_trends_integration_available": trends_configured,
        "voice_synthesis_integration_available": voice_configured,
        "image_generation_integration_available": image_configured,
        "prayer_notifications_available": prayer_configured,
        "location_detection_available": True,
        "file_processing_available": True,
        "note": "This is a helper module - endpoints are handled by router.py"
    }
#-- Section 13: Pattern Fatigue Detection Functions 9/29/25
def detect_pattern_complaint(message: str) -> tuple[bool, str, str]:
    """Detect if user is complaining about repetitive patterns"""
    message_lower = message.lower()
    
    duplicate_triggers = ["stop mentioning duplicate", "ignore duplicate", "stop pointing out double", "enough with the double", "stop saying twice", "stop being annoying", "quit that", "enough", "stop that"]
    time_joke_triggers = ["stop with the 2am", "enough 2am jokes", "quit asking about 2am", "stop time jokes", "no more 2am"]
    
    for trigger in duplicate_triggers:
        if trigger in message_lower:
            return True, "duplicate_callouts", message
    
    for trigger in time_joke_triggers:
        if trigger in message_lower:
            return True, "2am_jokes", message
    
    return False, "", ""

async def handle_pattern_complaint(user_id: str, pattern_type: str, complaint_text: str) -> str:
    """Handle user complaints about annoying patterns"""
    try:
        if pattern_type == "duplicate_callouts":
            return await handle_duplicate_complaint(user_id, complaint_text)
        elif pattern_type == "2am_jokes":
            return await handle_time_joke_complaint(user_id, complaint_text)
        else:
            return "Got it! I'll try to be less repetitive with that pattern."
    except Exception as e:
        logger.error(f"Failed to handle pattern complaint: {e}")
        return "I'll try to be less repetitive with that pattern."

def detect_pattern_fatigue_command(message: str) -> bool:
    """Detect pattern fatigue management commands"""
    return any(keyword in message.lower() for keyword in ["pattern fatigue", "pattern stats", "fatigue status"])

#-- Section 14: Google Workspace Integration Functions - 9/30/25
#-- Section 14: Google Workspace Integration Functions - 9/30/25
#-- Section 14: Google Workspace Integration Functions - 9/30/25
#-- Section 14: Google Workspace Integration Functions updated for web autho - 10/1/25
#-- Section 14: Google Workspace Integration Functions - OAuth web flow + Keywords/Analytics handlers
def detect_google_command(message: str) -> tuple[bool, str]:
    """Detect Google Workspace commands and determine command type"""
    google_keywords = [
            'google auth', 'google status', 'google sites', 'google accounts',
            'google keywords', 'google analytics', 'google drive', 'google email',
            'google gmail', 'google calendar', 'google suggest', 'google patterns',
            'google predict', 'google intelligence', 'google optimal',
            'copy to drive', 'save to google doc', 'move to drive',  # ← Already added
            'copy that to drive', 'save that to drive',  # ← ADD THESE
    ]
    
    message_lower = message.lower()
    
    # Check if it's a Google command
    is_google_command = any(keyword in message_lower for keyword in google_keywords)
    
    if not is_google_command:
        return False, ''
    
    # Determine command type
    if 'auth' in message_lower:
        if 'setup' in message_lower or 'start' in message_lower:
            return True, 'auth_setup'
        elif 'status' in message_lower:
            return True, 'auth_status'
        elif 'accounts' in message_lower:
            return True, 'auth_accounts'
        else:
            return True, 'auth_help'
    
    elif 'keywords' in message_lower or 'keyword' in message_lower:
        if 'pending' in message_lower:
            return True, 'keywords_pending'
        elif 'approve' in message_lower:
            return True, 'keywords_approve'
        elif 'ignore' in message_lower:
            return True, 'keywords_ignore'
        else:
            return True, 'keywords_view'
    
    elif 'analytics' in message_lower:
        if 'all' in message_lower:
            return True, 'analytics_all'
        else:
            return True, 'analytics_site'
    
    elif 'optimal' in message_lower and 'timing' in message_lower:
        return True, 'optimal_timing'
    
    elif 'drive' in message_lower:
            if 'copy' in message_lower or 'save' in message_lower or 'move' in message_lower:
                return True, 'drive_copy'
            elif 'create' in message_lower and 'doc' in message_lower:
                return True, 'drive_create_doc'
            elif 'create' in message_lower and 'sheet' in message_lower:
                return True, 'drive_create_sheet'
            elif 'recent' in message_lower:
                return True, 'drive_recent'
            else:
                return True, 'drive_help'
    
    elif 'email' in message_lower or 'gmail' in message_lower:
        if 'summary' in message_lower:
            return True, 'email_summary'
        elif 'draft' in message_lower:
            return True, 'email_draft'
        else:
            return True, 'email_account'
    
    elif 'calendar' in message_lower:
        if 'today' in message_lower:
            return True, 'calendar_today'
        elif 'week' in message_lower:
            return True, 'calendar_week'
        elif 'feeds' in message_lower:
            return True, 'calendar_feeds'
        elif 'windows' in message_lower:
            return True, 'calendar_windows'
        else:
            return True, 'calendar_help'
    
    elif 'suggest' in message_lower or 'intelligence' in message_lower:
        return True, 'intelligence_suggest'
    
    elif 'patterns' in message_lower:
        return True, 'intelligence_patterns'
    
    elif 'predict' in message_lower:
        return True, 'intelligence_predict'
    
    elif 'status' in message_lower:
        return True, 'status_check'
    
    elif 'sites' in message_lower:
        return True, 'sites_list'
    
    else:
        return True, 'general_help'

async def process_google_command(message: str, user_id: str, thread_id: Optional[str] = None) -> str:
    """Process Google Workspace commands and return personality-driven responses"""
    try:
        import httpx
        
        # ⚡ FIX: Get command_type by calling detect_google_command FIRST
        is_google_cmd, command_type = detect_google_command(message)
        
        if not is_google_cmd:
            return "Not a valid Google Workspace command. Try `google auth setup` or `google help`"
                
        # Get base URL for API calls
        base_url = "http://localhost:8000"
        
        async with httpx.AsyncClient(timeout=30.0) as client:
            
            # ============================================================
            # AUTHENTICATION COMMANDS
            # ============================================================
            
            if command_type == 'auth_setup':
                response = await client.get(f"{base_url}/google/auth/start?user_id={user_id}")
                data = response.json()
                
                if data.get('success'):
                    auth_url = data['authorization_url']
                    
                    return f"""🔐 **Google Workspace Authentication**

Click this link to authorize access:
{auth_url}

After you approve, you'll be redirected back to the chat.

This will give me access to:
- Google Analytics (read-only)
- Search Console (read-only)  
- Google Drive (create/edit documents)
- Gmail (read + compose)
- Your email and profile info

All tokens are encrypted and stored securely."""
                else:
                    return f"❌ **Authentication Failed**\n\n{data.get('message', 'Unknown error')}"
            
            elif command_type == 'auth_status':
                # Check OAuth status
                response = await client.get(f"{base_url}/google/auth/accounts?user_id={user_id}")
                data = response.json()
                
                if data.get('count', 0) > 0:
                    accounts = data['accounts']
                    account_list = "\n".join([f"   • {acc['email']} - {'✅ Active' if not acc.get('is_expired') else '⚠️ Expired'}" for acc in accounts])
                    return f"""Google Auth Status

Connected Accounts ({data['count']}):
{account_list}

Use `google auth setup` to add more accounts!"""
                else:
                    return "No Google accounts connected yet. Use `google auth setup` to connect!"
            
            elif command_type == 'auth_accounts':
                response = await client.get(f"{base_url}/google/auth/accounts")
                data = response.json()
                
                if data.get('count', 0) > 0:
                    accounts = data['accounts']
                    account_details = []
                    for acc in accounts:
                        status = "✅ Active" if not acc.get('is_expired') else '⚠️ Expired'
                        account_details.append(f"""Account: {acc['email']}
Status: {status}
Scopes: {len(acc.get('scopes', []))} permissions
Connected: {acc.get('authenticated_at', 'Unknown')}""")
                    
                    return f"""Connected Google Accounts

{chr(10).join(account_details)}

Total: {data['count']} account(s)"""
                else:
                    return "No Google accounts connected yet. Use `google auth setup` to connect!"
            
            # ============================================================
            # KEYWORDS/SEARCH CONSOLE COMMANDS
            # ============================================================
            
            elif command_type == 'keywords_view':
                # Extract site name from message
                site_name = None
                for site in ['bcdodge', 'rose_angel', 'meals_feelz', 'tv_signals', 'damn_it_carl']:
                    if site in message.lower() or site.replace('_', '') in message.lower():
                        site_name = site
                        break
                
                if not site_name:
                    return """Keyword Opportunities

Please specify which site you want to check:
- `google keywords bcdodge`
- `google keywords rose_angel`  
- `google keywords meals_feelz`
- `google keywords tv_signals`
- `google keywords damn_it_carl`"""
                
                try:
                    response = await client.get(f"{base_url}/google/keywords/opportunities?site_name={site_name}&user_id={user_id}")
                    data = response.json()
                    
                    opportunities = data.get('opportunities', [])
                    
                    if not opportunities:
                        return f"""Keyword Opportunities: {site_name}

No new keyword opportunities found for this site.

This means:
- All high-value keywords (100+ impressions, positions 11-30) are already tracked
- Keep creating content and check back later!

Use `google keywords pending` to see opportunities across all sites."""
                    
                    result = f"""Keyword Opportunities: {site_name}

Found {len(opportunities)} new opportunities:

"""
                    for i, opp in enumerate(opportunities[:10], 1):  # Show top 10
                        result += f"""{i}. **{opp['keyword']}**
   📊 Impressions: {opp['impressions']} | Clicks: {opp['clicks']}
   📍 Position: {opp['position']:.1f} | Type: {opp['opportunity_type']}
   🌐 Page: {opp.get('page', 'N/A')}
   🌍 Country: {opp.get('country', 'N/A')} | 📱 Device: {opp.get('device', 'N/A')}
   
"""
                    
                    result += "\nThese keywords have high impressions but aren't being tracked yet. Ready to add them?"
                    return result
                    
                except Exception as e:
                    logger.error(f"Keywords command failed: {e}")
                    return f"Error fetching keyword opportunities: {str(e)}"
            
            elif command_type == 'keywords_pending':
                return """Pending Keywords Across All Sites

This feature shows keyword opportunities across all your sites.

For now, check each site individually:
- `google keywords bcdodge`
- `google keywords rose_angel`
- `google keywords meals_feelz`
- `google keywords tv_signals`
- `google keywords damn_it_carl`"""
            
            # ============================================================
            # ANALYTICS COMMANDS
            # ============================================================
            
            elif command_type == 'analytics_site':
                # Extract site name from message
                site_name = None
                for site in ['bcdodge', 'rose_angel', 'meals_feelz', 'tv_signals', 'damn_it_carl']:
                    if site in message.lower() or site.replace('_', '') in message.lower():
                        site_name = site
                        break
                
                if not site_name:
                    return """Analytics Summary

Please specify which site you want to check:
- `google analytics bcdodge`
- `google analytics rose_angel`
- `google analytics meals_feelz`
- `google analytics tv_signals`
- `google analytics damn_it_carl`

Or use `google analytics all` for an overview."""
                
                try:
                    response = await client.get(f"{base_url}/google/analytics/summary?site_name={site_name}&user_id={user_id}")
                    data = response.json()
                    
                    stats = data.get('summary', {})
                    
                    # Build epic analytics display
                    result = f"""🎯 **Analytics Summary: {site_name}**
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📊 **TRAFFIC OVERVIEW** (Last 30 Days)

👥 **Visitors:** {stats.get('total_visitors', 0):,} total
   ├─ 🆕 New: {stats.get('new_users', 0):,} ({(stats.get('new_users', 0) / stats.get('total_visitors', 1) * 100):.1f}%)
   └─ 🔄 Returning: {stats.get('returning_users', 0):,} ({(stats.get('returning_users', 0) / stats.get('total_visitors', 1) * 100):.1f}%)

📄 **Page Views:** {stats.get('total_pageviews', 0):,}
💬 **Sessions:** {stats.get('total_sessions', 0):,}
⏱️ **Avg. Session:** {stats.get('avg_session_duration', 0):.1f}s
🎯 **Engagement Rate:** {stats.get('engagement_rate', 0):.1f}%
📊 **Bounce Rate:** {stats.get('bounce_rate', 0):.1f}%
"""

                    # Add device breakdown if available
                    devices = stats.get('devices', [])
                    if devices:
                        result += "\n📱 **DEVICES:**\n"
                        for device in devices[:3]:
                            device_name = device.get('device', 'unknown')
                            device_sessions = device.get('sessions', 0)
                            result += f"   • {device_name.title()}: {device_sessions:,} sessions\n"
                    
                    # Add top pages if available
                    top_pages = stats.get('top_pages', [])
                    if top_pages:
                        result += "\n🔥 **TOP PAGES:**\n"
                        for i, page in enumerate(top_pages[:5], 1):
                            result += f"   {i}. {page}\n"
                    
                    # Add traffic sources if available
                    traffic_sources = stats.get('traffic_sources', {})
                    if traffic_sources:
                        result += "\n🚀 **TRAFFIC SOURCES:**\n"
                        for source, count in list(traffic_sources.items())[:5]:
                            result += f"   • {source.title()}: {count:,}\n"
                    
                    result += "\n💡 Use `google analytics all` to compare across all sites!"
                    
                    return result
                    
                except Exception as e:
                    logger.error(f"Analytics command failed: {e}")
                    return f"Error fetching analytics: {str(e)}"
            
            elif command_type == 'analytics_all':
                return """Analytics Overview - All Sites

Feature coming soon! For now, check each site individually:
- `google analytics bcdodge`
- `google analytics rose_angel`
- `google analytics meals_feelz`
- `google analytics tv_signals`
- `google analytics damn_it_carl`"""
            
            # ============================================================
            # EMAIL/GMAIL COMMANDS
            # ============================================================
            
            elif command_type == 'email_summary':
                from ..integrations.google_workspace.gmail_client import get_email_summary
                
                try:
                    summary = await get_email_summary(user_id, days=7)
                    
                    if not summary or summary.get('total_emails', 0) == 0:
                        return """Email Summary

No emails found in the last 7 days, or Gmail is not yet connected.

Use `google auth setup` to connect your Gmail account."""
                    
                    # Build the base summary
                    response_parts = [
                        f"""Email Summary (Last 7 Days)

Total Emails: {summary.get('total_emails', 0)}
Important: {summary.get('important', 0)}
Needs Response: {summary.get('needs_response', 0)}
Inbox: {summary.get('inbox', 0)}
Sent: {summary.get('sent', 0)}
Negative Sentiment: {summary.get('negative_sentiment', 0)}"""
                    ]
                    
                    # Add emails requiring response if any exist
                    emails_needing_response = summary.get('emails_requiring_response', [])
                    if emails_needing_response:
                        response_parts.append("\n\n📧 **Emails Requiring Response:**\n")
                        
                        for i, email in enumerate(emails_needing_response[:10], 1):  # Show max 10
                            priority_emoji = "🔥" if email.get('priority') == 'high' else "📨"
                            response_parts.append(f"""{i}. {priority_emoji} **From:** {email['from']}
   **Subject:** {email['subject']}
   **Date:** {email['date'].strftime('%b %d, %I:%M %p') if hasattr(email['date'], 'strftime') else email['date']}
""")
                    
                    response_parts.append("\nUse `google email draft` to create responses!")
                    
                    return "".join(response_parts)
                    
                except Exception as e:
                    logger.error(f"Email summary failed: {e}")
                    return f"""Email Summary Error

Could not retrieve email summary: {str(e)}

Make sure your Gmail account is connected with `google auth setup`"""
            
            elif command_type == 'email_draft':
                return """Create Email Draft

To create an email draft, I need more information:

Example format:
"Create an email draft to john@example.com about meeting tomorrow"

Or provide:
- To: Recipient email
- Subject: Email subject
- Body: Your message

Once you provide these details, I'll create a draft in your Gmail account that you can review and send."""
            
            elif command_type == 'email_account':
                return """Gmail Account Management

Your Gmail integration provides:
- Email analysis and prioritization
- Smart filtering (urgent, business, needs response)
- Sentiment analysis
- Draft creation with AI assistance
- Privacy-first (metadata only, 30-day retention)

Commands:
- `google email summary` - See last 7 days
- `google email draft` - Create a draft

Connected Accounts: Check with `google auth accounts`"""
            
            # ============================================================
            # DRIVE COMMANDS
            # ============================================================
            
            elif command_type == 'drive_create_doc':
                # Parse title from message
                import re
                match = re.search(r'create doc(?:ument)?\s+(.+)', message.lower())
                if not match:
                    return """📄 **Create Google Doc**

Usage: `google drive create doc [title]`

Example: `google drive create doc Meeting Notes`

I'll create a new Google Doc with that title!"""
                
                title = match.group(1).strip()
                
                try:
                    from ..integrations.google_workspace.drive_client import create_google_doc
                    doc = await create_google_doc(user_id, title, "Document created from Syntax Prime chat")
                    
                    return f"""📄 **Document Created!**

**Title:** {doc['title']}
**Type:** Google Doc
**URL:** {doc['url']}

Your document is ready! Click the link above to open it in Google Docs."""
                    
                except Exception as e:
                    logger.error(f"Drive create doc failed: {e}")
                    return f"❌ **Drive Error:** {str(e)}\n\nMake sure you're connected with `google auth status`"
            
            elif command_type == 'drive_create_sheet':
                # Parse title from message
                import re
                match = re.search(r'create sheet\s+(.+)', message.lower())
                if not match:
                    return """📊 **Create Google Sheet**

Usage: `google drive create sheet [title]`

Example: `google drive create sheet Budget 2025`

I'll create a new Google Sheet with that title!"""
                
                title = match.group(1).strip()
                
                try:
                    from ..integrations.google_workspace.drive_client import create_google_sheet
                    sheet = await create_google_sheet(user_id, title, [["Created from Syntax Prime"]])
                    
                    return f"""📊 **Spreadsheet Created!**

**Title:** {sheet['title']}
**Type:** Google Sheet
**URL:** {sheet['url']}

Your spreadsheet is ready! Click the link above to open it in Google Sheets."""
                    
                except Exception as e:
                    logger.error(f"Drive create sheet failed: {e}")
                    return f"❌ **Drive Error:** {str(e)}\n\nMake sure you're connected with `google auth status`"
            
            elif command_type == 'drive_copy':
                # Parse title and message count
                import re
                
                # Extract title if provided (after "as")
                title_match = re.search(r'(?:as|titled?)\s+(.+?)(?:\s*$)', message, re.IGNORECASE)
                title = title_match.group(1).strip() if title_match else None
                
                # Extract number of messages if specified
                num_match = re.search(r'last\s+(\d+)\s+messages?', message.lower())
                num_messages = int(num_match.group(1)) if num_match else 1
                
                # If no title provided, ask for one
                if not title:
                    return """📄 **Copy to Google Drive**

            Please provide a title for the document:

            Usage: `copy that to drive as [title]`

            Examples:
            - `copy that to drive as Marketing Report`
            - `copy last 3 messages to drive as Meeting Notes`

            I'll grab the conversation content and create a formatted Google Doc!"""
                
                # Get conversation history
                try:
                    from ..ai.conversation_manager import get_memory_manager
                    memory = get_memory_manager(user_id)
                    
                    if not thread_id:
                        return "❌ No conversation thread found. Start a conversation first, then use `copy that to drive as [title]`"
                    
                    # Get recent messages from this thread
                    messages = await memory.get_conversation_history(thread_id, limit=num_messages * 2)  # Get extras to filter
                    
                    if not messages:
                        return "❌ No messages found in this conversation to copy."
                    
                    # Build document content from messages
                    content_parts = []
                    message_count = 0
                    
                    for msg in reversed(messages):  # Reverse to get chronological order
                        if message_count >= num_messages:
                            break
                            
                        role = msg.get('role', 'unknown')
                        text = msg.get('content', '')
                        
                        if role == 'user':
                            content_parts.append(f"**User:**\n{text}\n\n")
                            message_count += 1
                        elif role == 'assistant':
                            content_parts.append(f"**Assistant:**\n{text}\n\n")
                            message_count += 1
                    
                    if not content_parts:
                        return "❌ No conversation content found to copy."
                    
                    # Combine into document
                    document_content = f"# {title}\n\n" + "".join(content_parts)
                    
                    # Create the Google Doc with full markdown formatting
                    from ..integrations.google_workspace.drive_client import create_google_doc
                    doc = await create_google_doc(user_id, title, document_content, thread_id)
                    
                    messages_copied = "message" if num_messages == 1 else f"{num_messages} messages"
                    
                    return f"""📄 **Copied to Google Drive!**

            **Title:** {doc['title']}
            **Content:** Last {messages_copied}
            **URL:** {doc['url']}

            Your conversation has been saved as a Google Doc with full formatting!
            Click the link above to open and edit it."""
                    
                except Exception as e:
                    logger.error(f"Drive copy failed: {e}")
                    import traceback
                    logger.error(f"Traceback: {traceback.format_exc()}")
                    return f"❌ **Drive Copy Error:** {str(e)}\n\nMake sure you're connected with `google auth status`"
            
            elif command_type == 'drive_recent':
                try:
                    from ..integrations.google_workspace.drive_client import list_recent_docs
                    docs = await list_recent_docs(user_id, limit=10)
                    
                    if not docs:
                        return """📁 **Recent Google Drive Documents**

            No documents created yet through Syntax Prime.

            Use these commands to create your first document:
            - `google drive create doc [title]`
            - `google drive create sheet [title]`
            - `copy that to drive as [title]`"""
                    
                    response = f"📁 **Recent Google Drive Documents** ({len(docs)} documents)\n\n"
                    
                    for i, doc in enumerate(docs, 1):
                        doc_type = "📄 Doc" if doc['type'] == 'document' else "📊 Sheet"
                        created = doc['created'].strftime('%b %d, %I:%M %p') if hasattr(doc['created'], 'strftime') else str(doc['created'])
                        
                        response += f"{i}. {doc_type} **{doc['title']}**\n"
                        response += f"   Created: {created}\n"
                        response += f"   {doc['url']}\n\n"
                    
                    return response
                    
                except Exception as e:
                    logger.error(f"Drive recent failed: {e}")
                    return f"❌ **Drive Error:** {str(e)}\n\nMake sure you're connected with `google auth status`"
            
            elif command_type == 'drive_help':
                return """📁 **Google Drive Commands**

**Working Commands:**
   - `google drive create doc [title]` - Create new Google Doc
   - `google drive create sheet [title]` - Create new spreadsheet
   - `copy that to drive as [title]` - Save conversation to Drive
   - `copy last [N] messages to drive as [title]` - Save multiple messages
   - `google drive recent` - View recently created documents

Use `google auth status` to check your connection!"""

            # ============================================================
            # CALENDAR COMMANDS
            # ============================================================
            
            elif command_type == 'calendar_today':
                try:
                    from ..integrations.google_workspace.calendar_client import get_today_schedule
                    events = await get_today_schedule(user_id)
                    
                    if not events:
                        return "📅 **Today's Calendar**\n\nNo events scheduled for today. Enjoy your free time!"
                    
                    response = f"📅 **Today's Calendar** ({len(events)} events)\n\n"
                    for event in events:
                        start = event.get('start', {})
                        time_str = start.get('dateTime', start.get('date', 'Unknown'))[:16].replace('T', ' ')
                        cal_name = event.get('calendar_name', 'Unknown')
                        response += f"**{time_str}** - {event['summary']}\n"
                        response += f"   📍 {cal_name}\n\n"
                    
                    return response
                    
                except Exception as e:
                    logger.error(f"Calendar today failed: {e}")
                    return f"❌ **Calendar Error:** {str(e)}\n\nMake sure you're connected with `google auth status`"
            
            elif command_type == 'calendar_week':
                try:
                    from ..integrations.google_workspace.calendar_client import get_week_schedule
                    events = await get_week_schedule(user_id)
                    
                    if not events:
                        return "📅 **This Week's Calendar**\n\nNo events scheduled this week. Time to plan something!"
                    
                    # Group events by day
                    from datetime import datetime
                    events_by_day = {}
                    for event in events:
                        start = event.get('start', {})
                        date_str = start.get('dateTime', start.get('date', ''))[:10]
                        if date_str not in events_by_day:
                            events_by_day[date_str] = []
                        events_by_day[date_str].append(event)
                    
                    response = f"📅 **This Week's Calendar** ({len(events)} events)\n\n"
                    
                    for date_str in sorted(events_by_day.keys()):
                        day_events = events_by_day[date_str]
                        date_obj = datetime.fromisoformat(date_str)
                        day_name = date_obj.strftime('%A, %B %d')
                        
                        response += f"**{day_name}** ({len(day_events)} events)\n"
                        for event in day_events[:5]:  # Limit to 5 per day
                            start = event.get('start', {})
                            time_str = start.get('dateTime', start.get('date', ''))[:16].replace('T', ' at ')
                            response += f"  • {event['summary']} - {time_str}\n"
                        
                        if len(day_events) > 5:
                            response += f"  • ... and {len(day_events) - 5} more\n"
                        response += "\n"
                    
                    return response
                    
                except Exception as e:
                    logger.error(f"Calendar week failed: {e}")
                    return f"❌ **Calendar Error:** {str(e)}\n\nMake sure you're connected with `google auth status`"
            
            elif command_type == 'calendar_feeds':
                try:
                    from ..integrations.google_workspace.calendar_client import get_calendar_feeds
                    calendars = await get_calendar_feeds(user_id)
                    
                    if not calendars:
                        return "📅 **Calendar Feeds**\n\nNo calendars found. Check your Google Calendar connection."
                    
                    response = f"📅 **Your Calendar Feeds** ({len(calendars)} calendars)\n\n"
                    
                    for cal in calendars:
                        primary = " (Primary)" if cal.get('primary') else ""
                        selected = "✓" if cal.get('selected') else "○"
                        response += f"{selected} **{cal['summary']}**{primary}\n"
                        response += f"   Access: {cal.get('access_role', 'unknown')}\n\n"
                    
                    return response
                    
                except Exception as e:
                    logger.error(f"Calendar feeds failed: {e}")
                    return f"❌ **Calendar Error:** {str(e)}\n\nMake sure you're connected with `google auth status`"
            
            elif command_type == 'calendar_windows':
                return "📅 This feature hasn't been built yet. Use `google calendar week` to view your schedule."
            
            elif command_type == 'calendar_help':
                return """📅 **Google Calendar Commands**

**Working Commands:**
   - `google calendar today` - Today's schedule
   - `google calendar week` - This week's events
   - `google calendar feeds` - List all your calendars

Use `google auth status` to check your connection!"""
            
            # ============================================================
            # STATUS & HELP COMMANDS
            # ============================================================
            
            elif command_type == 'status_check':
                response = await client.get(f"{base_url}/google/status")
                data = response.json()
                
                health = "Healthy" if data.get('overall_health') == 'healthy' else "Issues Detected"
                oauth_count = len(data.get('oauth_authentication', {}).get('accounts', []))
                analytics_count = len(data.get('analytics_sites', []))
                
                return f"""Google Workspace Status

System Health: {health}
OAuth Accounts: {oauth_count} connected
Analytics Sites: {analytics_count} configured
Features Active:
   - Search Console keyword tracking
   - Analytics insights
   - Gmail intelligence (multi-account)
   - Calendar integration (11 feeds)
   - Self-evolving Intelligence Engine
   - Drive document creation

Use `google auth accounts` to see connected accounts."""
            
            elif command_type == 'sites_list':
                return """Configured Sites

Active Sites:
   1. bcdodge - Primary site
   2. rose_angel - Meditation & wellness
   3. meals_feelz - Food & recipes
   4. tv_signals - TV & entertainment
   5. damn_it_carl - Personal blog

Each site has Analytics and Search Console tracking enabled. Use `google keywords [site]` to see opportunities!"""
            
            else:
                return """Google Workspace Commands

Authentication:
   - `google auth setup` - Connect your Google accounts
   - `google auth status` - Check connection status
   - `google auth accounts` - List connected accounts

Keywords (Search Console):
   - `google keywords [site]` - View opportunities for a site
   - `google keywords pending` - See all pending across sites

Analytics:
   - `google analytics [site]` - Traffic summary
   - `google analytics all` - All sites overview

Email:
   - `google email summary` - Last 7 days of email
   - `google email draft` - Create email draft
   - `google email account` - Gmail integration info

Status:
   - `google status` - Integration health check
   - `google sites` - List configured sites

More features coming: Calendar, Drive, and Intelligence commands!"""
    
    except Exception as e:
        logger.error(f"Google command processing failed: {e}")
        return f"Google Workspace Error: {str(e)}\n\nUse `google auth setup` to connect!"

def get_integration_info() -> Dict[str, Any]:
    """Get Google Workspace integration information"""
    return {
        "name": "Google Workspace Integration",
        "version": "2.0.0",
        "status": "active",
        "features": [
            "OAuth web flow authentication",
            "Search Console keyword opportunities",
            "Google Analytics insights",
            "Gmail multi-account intelligence",
            "Drive document creation"
        ]
    }

def check_module_health() -> Dict[str, Any]:
    """Check Google Workspace module health"""
    return {
        "healthy": True,
        "module": "google_workspace",
        "status": "operational"
    }

#-- Section 14.5: Email Detail and Action Commands - 10/2/25
#-- Section 14.5: Email Detail and Action Commands - 10/2/25
def detect_email_detail_command(message: str) -> tuple[bool, str, Optional[int]]:
    """
    Detect email detail/action requests (e.g., "summarize email 4")
    
    Returns: (is_email_command, action_type, email_number)
    action_type: 'summarize', 'reply', 'read'
    
    Examples:
    - "summarize email 4" -> (True, 'summarize', 4)
    - "reply to email #2" -> (True, 'reply', 2)
    - "read email 3" -> (True, 'read', 3)
    """
    import re
    message_lower = message.lower()
    
    # Pattern 1: "summarize email 4" or "email #4 summary"
    summarize_patterns = [
        r'summarize email #?(\d+)',
        r'email #?(\d+) summary',
        r'what(?:\'s| is) email #?(\d+) about',
        r'tell me about email #?(\d+)',
        r'details (?:on|for|about) email #?(\d+)'
    ]
    
    for pattern in summarize_patterns:
        match = re.search(pattern, message_lower)
        if match:
            return True, 'summarize', int(match.group(1))
    
    # Pattern 2: "reply to email 4" or "draft reply to #4"
    reply_patterns = [
        r'reply to email #?(\d+)',
        r'draft (?:a )?reply (?:to|for) #?(\d+)',
        r'respond to email #?(\d+)',
        r'answer email #?(\d+)'
    ]
    
    for pattern in reply_patterns:
        match = re.search(pattern, message_lower)
        if match:
            return True, 'reply', int(match.group(1))
    
    # Pattern 3: "read email 4" or "show me email #4"
    read_patterns = [
        r'read email #?(\d+)',
        r'show (?:me )?email #?(\d+)',
        r'open email #?(\d+)',
        r'view email #?(\d+)'
    ]
    
    for pattern in read_patterns:
        match = re.search(pattern, message_lower)
        if match:
            return True, 'read', int(match.group(1))
    
    return False, '', None

async def process_email_detail_command(action_type: str, email_index: int, user_id: str) -> str:
    """
    Process email detail/action commands
    
    Args:
        action_type: 'summarize', 'reply', or 'read'
        email_index: Email number from summary list (1-10)
        user_id: User ID for authentication
        
    Returns:
        Formatted response with email details or action results
    """
    try:
        from ..integrations.google_workspace.gmail_client import gmail_client
        
        # Initialize if needed
        if not gmail_client._user_id:
            await gmail_client.initialize(user_id)
        
        # SUMMARIZE: Show key points and preview
        if action_type == 'summarize':
            return await gmail_client.summarize_email(email_index)
        
        # REPLY: Prompt for reply content
        # REPLY: Generate AI suggestion AND create draft
        elif action_type == 'reply':
            email = await gmail_client.get_email_by_index(email_index)
            if not email:
                return f"Email #{email_index} not found. Run `google email summary` first."
            
            # Extract sender email from "Name <email>" format
            import re
            sender_match = re.search(r'<(.+?)>', email['from'])
            sender_email = sender_match.group(1) if sender_match else email['from']
            
            # Generate AI-suggested reply using OpenRouter
            from ..ai.openrouter_client import get_openrouter_client
            openrouter = await get_openrouter_client()
            
            ai_response = await openrouter.chat_completion(
                messages=[{
                    "role": "system",
                    "content": f"Generate a professional email reply to this message. Keep it concise and appropriate. Original email:\n\nFrom: {email['from']}\nSubject: {email['subject']}\nDate: {email['date']}\n\nBody:\n{email['body'][:1000]}"
                }, {
                    "role": "user",
                    "content": "Generate a suggested reply to this email."
                }],
                model="anthropic/claude-3.5-sonnet:beta",
                max_tokens=500
            )
            
            suggested_reply = ai_response['choices'][0]['message']['content']
            
            # Create the draft
            try:
                draft_result = await gmail_client.create_draft(
                    email=None,
                    to=sender_email,
                    subject=f"Re: {email['subject']}",
                    body=suggested_reply
                )
                
                return f"""✅ **Draft Reply Created for Email #{email_index}**

        **To:** {sender_email}
        **Subject:** Re: {email['subject']}

        **Draft Content:**
        {suggested_reply}

        ---
        **Draft ID:** {draft_result['id']}

        The draft has been saved to your Gmail drafts. You can review and edit it before sending!

        Commands:
        - Open Gmail to send the draft
        - `reply to email {email_index}` to regenerate with different content"""
            
            except Exception as draft_error:
                logger.error(f"Failed to create draft: {draft_error}")
                return f"""**Suggested Reply for Email #{email_index}**

        **To:** {sender_email}
        **Subject:** Re: {email['subject']}

        {suggested_reply}

        ---
        ⚠️ Could not save to Gmail drafts: {str(draft_error)}

        You can copy this text and create the draft manually."""
        
        # READ: Show full email body
        elif action_type == 'read':
            email = await gmail_client.get_email_by_index(email_index)
            if not email:
                return f"Email #{email_index} not found. Run `google email summary` first."
            
            # Truncate very long bodies
            body_display = email['body']
            if len(body_display) > 3000:
                body_display = body_display[:3000] + "\n\n[... body truncated, showing first 3000 characters ...]"
            
            return f"""Full Email #{email_index}

From: {email['from']}
To: {email['to']}
Subject: {email['subject']}
Date: {email['date']}

Body:
{body_display}

---
Actions: 
- `reply to email {email_index}` - Draft a response
- `summarize email {email_index}` - Get key points"""
        
        else:
            return f"Unknown email action type: {action_type}"
        
    except Exception as e:
        logger.error(f"Email detail command failed: {e}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        return f"Error processing email command: {str(e)}"

def detect_draft_creation_command(message: str) -> tuple[bool, Optional[int], Optional[str]]:
    """
    Detect draft creation requests
    
    Returns: (is_draft_cmd, email_reply_number, draft_instruction)
    
    Examples:
    - "save as draft" -> (True, None, None) - uses last AI response
    - "create draft replying thanks" -> (True, None, "thanks")
    - "save this as a draft" -> (True, None, None)
    """
    import re
    message_lower = message.lower()
    
    # Patterns that indicate wanting to save a draft
    draft_patterns = [
        r'save (?:this |that )?(?:as |to )?draft',
        r'create (?:a )?draft',
        r'make (?:this |that )?(?:a )?draft',
        r'save to drafts',
        r'draft this'
    ]
    
    for pattern in draft_patterns:
        if re.search(pattern, message_lower):
            # Check if they're providing additional instruction
            instruction_patterns = [
                r'(?:saying|with|that says?|about)\s+(.+)',
                r'draft:\s*(.+)',
                r'draft\s+(.+)'
            ]
            
            for inst_pattern in instruction_patterns:
                match = re.search(inst_pattern, message_lower)
                if match:
                    return True, None, match.group(1).strip()
            
            return True, None, None
    
    return False, None, None

async def process_draft_creation_command(conversation_history: List[Dict], user_id: str, custom_instruction: Optional[str] = None) -> str:
    """
    Create a Gmail draft from conversation context
    
    Args:
        conversation_history: Recent messages to extract draft content from
        user_id: User ID for Gmail authentication
        custom_instruction: Optional custom text for the draft
        
    Returns:
        Confirmation message with draft details
    """
    try:
        from ..integrations.google_workspace.gmail_client import gmail_client
        
        # Initialize if needed
        if not gmail_client._user_id:
            await gmail_client.initialize(user_id)
        
        # Extract draft content from conversation
        # Look for the last assistant message that looks like an email
        draft_content = None
        recipient = None
        subject = None
        
        # Search backwards through conversation for email content
        for msg in reversed(conversation_history):
            if msg['role'] == 'assistant':
                content = msg['content']
                
                # Try to parse email structure
                if 'Subject:' in content or 'To:' in content:
                    # Parse structured email
                    lines = content.split('\n')
                    body_lines = []
                    
                    for line in lines:
                        if line.startswith('To:'):
                            recipient = line.replace('To:', '').strip()
                        elif line.startswith('Subject:') or line.startswith('Re:'):
                            subject = line.replace('Subject:', '').strip()
                        elif line.strip() and not line.startswith(('From:', 'Date:', '**', '---', 'Actions:')):
                            body_lines.append(line)
                    
                    if body_lines:
                        draft_content = '\n'.join(body_lines).strip()
                        break
        
        # Use custom instruction if provided
        if custom_instruction:
            draft_content = custom_instruction
        
        if not draft_content:
            return """Unable to Create Draft

I couldn't find draft content in our recent conversation. 

To create a draft, either:
1. Generate a draft response first (e.g., "draft a reply to email 4")
2. Then say "save as draft"

Or provide the content directly:
- "create draft saying [your message here]"
- "save as draft: [your message here]"

What would you like to do?"""
        
        # Get recipient and subject from context if not found
        if not recipient:
            # Try to extract from earlier in conversation
            for msg in reversed(conversation_history):
                if 'email' in msg['content'].lower():
                    # Look for email addresses
                    import re
                    emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', msg['content'])
                    if emails:
                        recipient = emails[0]
                        break
            
            if not recipient:
                recipient = "recipient@example.com"  # Placeholder
        
        if not subject:
            subject = "Draft Reply"
        
        # Create the draft
        try:
            draft_result = await gmail_client.create_draft(
                email=None,  # Use default account
                to=recipient,
                subject=subject,
                body=draft_content
            )
            
            return f"""Draft Created Successfully

**Draft ID:** {draft_result['id']}
**To:** {recipient}
**Subject:** {subject}

**Content Preview:**
{draft_content[:200]}{'...' if len(draft_content) > 200 else ''}

The draft has been saved to your Gmail drafts folder. You can:
- Open Gmail to review and send it
- Edit it before sending
- Delete it if you change your mind

Draft is ready to send whenever you are!"""
            
        except Exception as e:
            logger.error(f"Failed to create Gmail draft: {e}")
            return f"""Draft Creation Failed

Could not save draft to Gmail: {str(e)}

**Draft Content (for your reference):**
To: {recipient}
Subject: {subject}

{draft_content}

You can copy this and create the draft manually in Gmail."""
        
    except Exception as e:
        logger.error(f"Draft creation command processing failed: {e}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        return f"Error creating draft: {str(e)}"

#-- Section 15: Reminder Functions - 10/16/25
def detect_reminder_command(message: str) -> Optional[str]:
    """
    Detect reminder-related commands
    
    Returns:
        'create' - Create new reminder
        'list' - List all reminders
        'cancel_one' - Cancel specific reminder
        'cancel_all' - Kill all reminders
        None - Not a reminder command
    """
    message_lower = message.lower()
    
    # KILL SWITCH - Check this FIRST (most important!)
    kill_phrases = [
        "cancel all reminders",
        "delete all reminders",
        "kill all reminders",
        "stop all reminders",
        "remove all reminders",
        "clear all reminders",
        "kill reminders"
    ]
    if any(phrase in message_lower for phrase in kill_phrases):
        return 'cancel_all'
    
    # Cancel specific reminder
    cancel_keywords = ["cancel reminder", "delete reminder", "remove reminder"]
    if any(keyword in message_lower for keyword in cancel_keywords):
        return 'cancel_one'
    
    # List reminders
    list_keywords = [
        "list reminders", "show reminders", "my reminders",
        "list my reminders", "show my reminders",
        "what reminders", "view reminders"
    ]
    if any(keyword in message_lower for keyword in list_keywords):
        return 'list'
    
    # Create reminder (default if contains remind keywords)
    create_keywords = [
        "remind", "reminder", "remind me", "set reminder",
        "create reminder", "make reminder"
    ]
    if any(keyword in message_lower for keyword in create_keywords):
        return 'create'
    
    return None


async def process_reminder_create(message: str, user_id: str) -> str:
    """
    Create a new reminder from natural language
    
    Examples:
        - "remind me to call mom in 30 minutes"
        - "remind me about meeting tomorrow at 2pm"
        - "set reminder for Friday at 3pm"
    """
    try:
        from ..integrations.telegram.notification_types.reminder_notifications import (
            get_reminder_notification_handler
        )
        
        handler = get_reminder_notification_handler()
        result = await handler.create_reminder_from_text(
            reminder_text=message,
            original_message=message
        )
        
        if result['success']:
            scheduled_time = result['scheduled_for']
            reminder_text = result['reminder_text']
            
            # Convert UTC to Eastern for display
            import pytz
            eastern = pytz.timezone('America/New_York')
            if scheduled_time.tzinfo is not None:
                local_time = scheduled_time.astimezone(eastern)
            else:
                local_time = scheduled_time
            
            # Format time nicely
            time_str = local_time.strftime('%A, %B %d at %I:%M %p')
            reminder_id = result['reminder_id']
            
            return f"""⏰ **Reminder Created Successfully!**

📝 **What:** {reminder_text}
🕐 **When:** {time_str}
🆔 **ID:** `{reminder_id[:8]}...`

✅ You'll receive a Telegram notification at the scheduled time.

**Manage Reminders:**
- `list reminders` - See all pending reminders
- `cancel reminder {reminder_id[:8]}` - Cancel this specific reminder
- `cancel all reminders` - Delete all reminders (kill switch)
"""
        else:
            error_msg = result.get('error', 'Unknown error')
            suggestion = result.get('suggestion', '')
            
            return f"""❌ **Couldn't Parse Reminder Time**

**Error:** {error_msg}

{suggestion}

**Try these formats:**
- "remind me to [task] in 30 minutes"
- "remind me about [event] tomorrow at 2pm"
- "set reminder for [thing] on Friday at 3pm"
- "remind me to [action] next Monday at 9am"
"""
    
    except Exception as e:
        logger.error(f"Reminder creation failed: {e}")
        return f"""❌ **Reminder System Error**

Error: {str(e)}

Try: "remind me to [task] [when]"
"""


async def process_reminder_list(user_id: str) -> str:
    """List all pending reminders for the user"""
    try:
        from ..integrations.telegram.database_manager import get_telegram_db_manager
        
        db_manager = get_telegram_db_manager()
        reminders = await db_manager.get_all_reminders(user_id)
        
        if not reminders:
            return """📋 **No Pending Reminders**

You don't have any reminders scheduled.

Create one with: "remind me to [task] [when]"
"""
        
        # Build the list
        reminder_list = []
        for idx, reminder in enumerate(reminders, 1):
            reminder_id = str(reminder['id'])
            text = reminder['reminder_text']
            scheduled = reminder['scheduled_for']
            
            # Format time
            time_str = scheduled.strftime('%a, %b %d at %I:%M %p')
            
            # Calculate time until
            now = datetime.now()
            if scheduled.tzinfo is not None:
                from datetime import timezone
                now = datetime.now(timezone.utc)
            
            delta = scheduled - now
            
            if delta.total_seconds() < 0:
                time_until = "⚠️ Overdue"
            elif delta.days > 0:
                time_until = f"in {delta.days} day(s)"
            else:
                hours = delta.seconds // 3600
                minutes = (delta.seconds % 3600) // 60
                if hours > 0:
                    time_until = f"in {hours}h {minutes}m"
                else:
                    time_until = f"in {minutes}m"
            
            reminder_list.append(
                f"{idx}. **{text}**\n"
                f"   🕐 {time_str} ({time_until})\n"
                f"   🆔 `{reminder_id[:8]}...`"
            )
        
        reminders_text = "\n\n".join(reminder_list)
        
        return f"""📋 **Your Pending Reminders** ({len(reminders)} total)

{reminders_text}

**Commands:**
- `cancel reminder [id]` - Cancel a specific reminder
- `cancel all reminders` - 🔥 Delete all reminders (kill switch)
"""
    
    except Exception as e:
        logger.error(f"Failed to list reminders: {e}")
        return f"❌ **Error:** {str(e)}"


async def process_reminder_cancel(message: str, user_id: str) -> str:
    """
    Cancel reminder(s) - handles both single and mass deletion
    
    Examples:
        - "cancel reminder abc-123-456"
        - "cancel all reminders"
    """
    try:
        from ..integrations.telegram.database_manager import get_telegram_db_manager
        
        db_manager = get_telegram_db_manager()
        message_lower = message.lower()
        
        # Check if it's the KILL SWITCH
        kill_phrases = [
            "cancel all", "delete all", "kill all",
            "stop all", "remove all", "clear all"
        ]
        
        is_kill_switch = any(phrase in message_lower for phrase in kill_phrases)
        
        if is_kill_switch:
            # KILL ALL REMINDERS
            count = await db_manager.delete_all_reminders(user_id)
            
            if count > 0:
                return f"""🔥 **KILL SWITCH ACTIVATED**

✅ Successfully deleted **{count} reminder(s)**

All pending reminders have been cancelled. No more notifications will be sent.

Create new reminders with: "remind me to [task] [when]"
"""
            else:
                return """📋 **No Reminders to Cancel**

You don't have any pending reminders.

Create one with: "remind me to [task] [when]"
"""
        
        else:
            # Cancel specific reminder - extract ID
            import re
            
            # Try to find UUID or short ID in message
            uuid_pattern = r'[0-9a-f]{8}[-]?[0-9a-f]{4}[-]?[0-9a-f]{4}[-]?[0-9a-f]{4}[-]?[0-9a-f]{12}'
            short_pattern = r'[0-9a-f]{8}'
            
            uuid_match = re.search(uuid_pattern, message_lower)
            if uuid_match:
                reminder_id = uuid_match.group(0)
            else:
                short_match = re.search(short_pattern, message_lower)
                if short_match:
                    # Get short ID and find full ID from database
                    short_id = short_match.group(0)
                    reminders = await db_manager.get_all_reminders(user_id)
                    
                    matching = [r for r in reminders if str(r['id']).startswith(short_id)]
                    
                    if len(matching) == 1:
                        reminder_id = str(matching[0]['id'])
                    elif len(matching) > 1:
                        return f"""❌ **Ambiguous ID**

The ID `{short_id}` matches multiple reminders. Please use a longer ID:

{chr(10).join([f"- `{str(r['id'])[:12]}...` - {r['reminder_text']}" for r in matching])}
"""
                    else:
                        return f"""❌ **Reminder Not Found**

No reminder found with ID starting with `{short_id}`.

Use `list reminders` to see all IDs.
"""
                else:
                    return """❌ **No Reminder ID Found**

Please specify which reminder to cancel:

**Usage:**
- `cancel reminder [id]` - where [id] is from `list reminders`
- `cancel all reminders` - to delete all reminders

Example: `cancel reminder abc12345`
"""
            
            # Delete the reminder
            success = await db_manager.delete_reminder(reminder_id, user_id)
            
            if success:
                return f"""✅ **Reminder Cancelled**

Reminder `{reminder_id[:8]}...` has been deleted.

Use `list reminders` to see remaining reminders.
"""
            else:
                return f"""❌ **Could Not Cancel Reminder**

Reminder `{reminder_id[:8]}...` not found or already sent.

Use `list reminders` to see your pending reminders.
"""
    
    except Exception as e:
        logger.error(f"Failed to cancel reminder: {e}")
        return f"❌ **Error:** {str(e)}"

#-- Section 16: Fathom Meeting Integration Functions - 10/19/25
def detect_meeting_query(message: str) -> tuple[bool, str]:
    """
    Detect meeting-related queries with improved sensitivity
    
    Returns:
        tuple: (is_meeting_query, query_type)
        query_type: 'specific', 'recent', 'search', 'action_items'
    """
    # Expanded keywords for better detection
    meeting_keywords = [
        "meeting", "meetings", "discussed", "talked about",
        "action item", "action items", "follow up", "follow-up",
        "what did we decide", "what was decided", "meeting notes",
        "decompress", "debrief", "recap", "call", "calls",
        "conversation", "session", "discussion", "chat with",
        "spoke with", "talked with", "mentioned", "said in",
        "brought up", "covered", "went over", "reviewed"
    ]
    
    # Temporal indicators that suggest meeting context
    temporal_keywords = [
        "yesterday", "today", "last week", "this week",
        "last month", "monday", "tuesday", "wednesday",
        "thursday", "friday", "saturday", "sunday",
        "morning", "afternoon", "ago", "recent"
    ]
    
    # Date patterns (October 15, Oct 15, 10/15, etc)
    import re
    date_patterns = [
        r'\b(january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{1,2}\b',
        r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\s+\d{1,2}\b',
        r'\b\d{1,2}/\d{1,2}\b',
        r'\b\d{1,2}-\d{1,2}\b'
    ]
    
    message_lower = message.lower()
    
    # Check for date patterns
    has_date = any(re.search(pattern, message_lower, re.IGNORECASE) for pattern in date_patterns)
    
    # Check for temporal keywords
    has_temporal = any(word in message_lower for word in temporal_keywords)
    
    # Check for meeting keywords
    has_meeting_keyword = any(keyword in message_lower for keyword in meeting_keywords)
    
    # NEW LOGIC: If there's a date OR temporal reference, assume it might be about a meeting
    is_meeting = has_meeting_keyword or has_date or has_temporal
    
    if not is_meeting:
        return (False, None)
    
    # Determine query type
    if 'action item' in message_lower:
        return (True, 'action_items')
    elif any(word in message_lower for word in ['recent', 'latest', 'last', 'today', 'yesterday', 'this week']):
        return (True, 'recent')
    elif any(word in message_lower for word in ['search', 'find', 'look for']):
        return (True, 'search')
    else:
        return (True, 'specific')

async def search_meetings(
    query: str,
    user_id: str,
    query_type: str = 'recent',
    limit: int = 5
) -> str:
    """
    Search meetings and return formatted context for AI
    
    Args:
        query: User's search query
        user_id: User ID
        query_type: Type of search ('recent', 'specific', 'search', 'action_items')
        limit: Max number of meetings to return
    
    Returns:
        Formatted meeting context string
    """
    try:
        from modules.core.database import db_manager
        import re
        from datetime import datetime, timedelta
        
        # Build appropriate query based on type
        if query_type == 'recent':
            sql = """
                SELECT 
                    title,
                    meeting_date,
                    duration_minutes,
                    participants,
                    ai_summary,
                    key_points
                FROM fathom_meetings
                ORDER BY meeting_date DESC
                LIMIT $1
            """
            meetings = await db_manager.fetch_all(sql, limit)
        
        elif query_type == 'action_items':
            sql = """
                SELECT DISTINCT
                    m.title,
                    m.meeting_date,
                    m.ai_summary
                FROM fathom_meetings m
                WHERE m.ai_summary IS NOT NULL
                ORDER BY m.meeting_date DESC
                LIMIT $1
            """
            meetings = await db_manager.fetch_all(sql, limit)
        
        else:  # specific or search
            # NEW: Try to extract dates from the query
            meeting_date = None
            query_lower = query.lower()
            
            # Try to parse dates like "October 15", "Oct 15", etc.
            date_patterns = {
                r'(january|jan)\s+(\d{1,2})': (1, None),
                r'(february|feb)\s+(\d{1,2})': (2, None),
                r'(march|mar)\s+(\d{1,2})': (3, None),
                r'(april|apr)\s+(\d{1,2})': (4, None),
                r'(may)\s+(\d{1,2})': (5, None),
                r'(june|jun)\s+(\d{1,2})': (6, None),
                r'(july|jul)\s+(\d{1,2})': (7, None),
                r'(august|aug)\s+(\d{1,2})': (8, None),
                r'(september|sep|sept)\s+(\d{1,2})': (9, None),
                r'(october|oct)\s+(\d{1,2})': (10, None),
                r'(november|nov)\s+(\d{1,2})': (11, None),
                r'(december|dec)\s+(\d{1,2})': (12, None),
            }
            
            for pattern, (month, _) in date_patterns.items():
                match = re.search(pattern, query_lower, re.IGNORECASE)
                if match:
                    day = int(match.group(2))
                    # Assume current year or last year if month hasn't happened yet this year
                    current_year = datetime.now().year
                    try:
                        meeting_date = datetime(current_year, month, day)
                        # If the date is in the future, try last year
                        if meeting_date > datetime.now():
                            meeting_date = datetime(current_year - 1, month, day)
                    except ValueError:
                        # Invalid date, skip
                        pass
                    break
            
            # If we found a date, search by date range (±3 days)
            if meeting_date:
                start_date = meeting_date - timedelta(days=3)
                end_date = meeting_date + timedelta(days=3)
                
                sql = """
                    SELECT 
                        title,
                        meeting_date,
                        duration_minutes,
                        participants,
                        ai_summary,
                        key_points
                    FROM fathom_meetings
                    WHERE meeting_date BETWEEN $1 AND $2
                    ORDER BY ABS(EXTRACT(EPOCH FROM (meeting_date - $3)))
                    LIMIT $4
                """
                meetings = await db_manager.fetch_all(sql, start_date, end_date, meeting_date, limit)
            else:
                # Fallback: Simple text search in title and summary
                search_term = f"%{query.lower()}%"
                sql = """
                    SELECT 
                        title,
                        meeting_date,
                        duration_minutes,
                        participants,
                        ai_summary,
                        key_points
                    FROM fathom_meetings
                    WHERE 
                        LOWER(title) LIKE $1
                        OR LOWER(ai_summary) LIKE $1
                    ORDER BY meeting_date DESC
                    LIMIT $2
                """
                meetings = await db_manager.fetch_all(sql, search_term, limit)
        
        if not meetings or len(meetings) == 0:
            return "\n\n📅 **Meeting Context:** No meetings found matching your query."
        
        # Format meetings for AI context
        if meeting_date:
            # If we searched by date, make it VERY clear we found meetings on/near that date
            meeting_context = [f"\n\n📅 **Meetings Found Near {meeting_date.strftime('%B %d, %Y')}:**\n"]
        else:
            meeting_context = ["\n\n📅 **Meeting Search Results:**\n"]
        
        for meeting in meetings:
            meeting_text = [
                f"\n**{meeting['title']}**",
                f"📅 Date: {meeting['meeting_date']}",
                f"⏱️ Duration: {meeting['duration_minutes']} minutes"
            ]
            
            if meeting.get('participants'):
                participants = meeting['participants']
                if isinstance(participants, list) and len(participants) > 0:
                    meeting_text.append(f"👥 Participants: {', '.join([str(p) for p in participants[:5]])}")
            
            if meeting.get('ai_summary'):
                # Show MORE of the summary so AI can make better matches
                summary = meeting['ai_summary']
                meeting_text.append(f"\n📝 Summary: {summary}")
            
            if meeting.get('key_points'):
                key_points = meeting['key_points']
                if isinstance(key_points, list) and len(key_points) > 0:
                    meeting_text.append("\n🔑 Key Points:")
                    for point in key_points[:5]:  # Show more key points
                        meeting_text.append(f"  • {point}")
            
            meeting_context.append("\n".join(meeting_text))
            meeting_context.append("---")
        
        context_text = "\n".join(meeting_context)
        
        # Make it VERY clear this is meeting data, not old knowledge
        if meeting_date:
            instruction = f"""
CRITICAL INSTRUCTION: The user asked about a meeting on or around {meeting_date.strftime('%B %d, %Y')}.
The meetings below were found within ±3 days of that date. Use these meetings to answer
the user's question. If the user mentioned a topic (like 'HubSpot'), search the summaries
below for that topic and use that meeting in your answer.
"""
        else:
            instruction = """
CRITICAL INSTRUCTION: The user is asking about RECENT MEETINGS from their Fathom
recording system. The meeting data below is from October 2025 and is MORE RECENT
and MORE ACCURATE than any older project information in your knowledge base.

When answering about meetings, you MUST use ONLY the information below.
DO NOT reference old project discussions from 2024 (Afghanistan, Dr. C, etc.).
"""
        
        return f"""

{'='*80}
🎯 FATHOM MEETING DATABASE - USE THIS DATA TO ANSWER
{'='*80}

{instruction}

{'='*80}

{context_text}

{'='*80}
END OF MEETING DATABASE - USE ONLY THIS FOR MEETING QUESTIONS
{'='*80}
"""
    
    except Exception as e:
        logger.error(f"Failed to search meetings: {e}")
        return "\n\n📅 **Meeting Context:** Unable to retrieve meeting information."

async def get_recent_meetings_context(user_id: str, days: int = 7, limit: int = 5) -> str:
    """
    Get recent meetings from the last N days - always available context
    
    This function is called proactively to ensure meetings are available
    even when the query doesn't explicitly mention them.
    
    Args:
        user_id: User ID
        days: How many days back to search (default 7)
        limit: Max number of meetings to return
    
    Returns:
        Formatted meeting context string or empty string if no meetings
    """
    try:
        from modules.core.database import db_manager
        from datetime import datetime, timedelta
        
        # Calculate date range
        end_date = datetime.now()
        start_date = end_date - timedelta(days=days)
        
        sql = """
            SELECT 
                title,
                meeting_date,
                duration_minutes,
                participants,
                ai_summary,
                key_points
            FROM fathom_meetings
            WHERE meeting_date >= $1
            ORDER BY meeting_date DESC
            LIMIT $2
        """
        
        meetings = await db_manager.fetch_all(sql, start_date, limit)
        
        if not meetings or len(meetings) == 0:
            logger.info(f"📅 No recent meetings found in last {days} days")
            return ""
        
        logger.info(f"📅 Found {len(meetings)} meetings from last {days} days")
        
        # Format meetings for AI context
        meeting_context = [f"\n\n📅 **Available Meeting Context (Last {days} Days):**\n"]
        
        for meeting in meetings:
            meeting_text = [
                f"\n**{meeting['title']}**",
                f"Date: {meeting['meeting_date']}",
                f"Duration: {meeting['duration_minutes']} minutes"
            ]
            
            if meeting.get('participants'):
                participants = meeting['participants']
                if isinstance(participants, list) and len(participants) > 0:
                    meeting_text.append(f"Participants: {', '.join([str(p) for p in participants[:5]])}")
            
            if meeting.get('ai_summary'):
                # Truncate long summaries for context
                summary = meeting['ai_summary']
                if len(summary) > 500:
                    summary = summary[:500] + "..."
                meeting_text.append(f"\nSummary: {summary}")
            
            if meeting.get('key_points'):
                key_points = meeting['key_points']
                if isinstance(key_points, list) and len(key_points) > 0:
                    meeting_text.append("\nKey Points:")
                    for point in key_points[:3]:  # Limit to 3 points
                        meeting_text.append(f"  • {point}")
            
            meeting_context.append("\n".join(meeting_text))
            meeting_context.append("---")
        
        context_text = "\n".join(meeting_context)
        
        return f"""

{'='*80}
🎯 RECENT MEETING CONTEXT - AVAILABLE FOR REFERENCE
{'='*80}

The following meetings occurred in the last {days} days. Use this information
to answer questions about recent discussions, decisions, or action items.

{'='*80}

{context_text}

{'='*80}
END OF RECENT MEETING CONTEXT
{'='*80}
"""
    
    except Exception as e:
        logger.error(f"❌ Failed to get recent meetings context: {e}")
        return ""
